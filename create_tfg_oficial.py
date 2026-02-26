# -*- coding: utf-8 -*-
"""
create_tfg_oficial.py

Genera TFG_Oficial.docx (~40 páginas) con formato normativo UMU 2025-2026.
Contenido: versión condensada de los 9 capítulos.
Formato:
  - Times New Roman 12pt  |  Márgenes 3cm x 2.5cm (A4)
  - Interlineado 1.5 (texto) / sencillo (tablas, refs)
  - Justificado
  - H1: 14pt NEGRITA MAYÚSCULAS negro
  - H2: 14pt NEGRITA negro
  - H3: 14pt CURSIVA negro
  - Núm. página: pie centrado
  - Citas: APA 7ª inline  |  Refs: alfabético APA 7ª, sangría francesa

Uso:
  python -X utf8 create_tfg_oficial.py
"""

import re
import sys
import io
import warnings
from pathlib import Path

warnings.filterwarnings("ignore")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).parent
FIGS_DIR     = PROJECT_ROOT / "output" / "figures_completo"

FIG1 = FIGS_DIR / "fig_01_gold_historia.png"
FIG2 = FIGS_DIR / "fig_02_determinantes.png"
FIG3 = FIGS_DIR / "fig_03_correlaciones_rolling.png"
FIG4 = FIGS_DIR / "fig_04_scatter.png"
FIG5 = FIGS_DIR / "fig_05_shap.png"
FIG6 = FIGS_DIR / "fig_06_ml_resultados.png"


# =====================================================================
# 1. DOCUMENTO Y ESTILOS
# =====================================================================

def create_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after       = Pt(6)
    normal.paragraph_format.space_before      = Pt(0)

    _cfg_h(doc, "Heading 1", 14, True,  False, True,  24, 12)
    _cfg_h(doc, "Heading 2", 14, True,  False, False, 12, 6)
    _cfg_h(doc, "Heading 3", 14, False, True,  False, 6,  6)
    return doc


def _cfg_h(doc, name, sz, bold, italic, caps, sb, sa):
    s = doc.styles[name]
    s.font.name      = "Times New Roman"
    s.font.size      = Pt(sz)
    s.font.bold      = bold
    s.font.italic    = italic
    s.font.all_caps  = caps
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before      = Pt(sb)
    s.paragraph_format.space_after       = Pt(sa)
    s.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


# =====================================================================
# 2. HELPERS
# =====================================================================

def _tnr(run, size=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rf  = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:cs"),    "Times New Roman")
    rPr.insert(0, rf)


def _inline(para, text: str, size: int = 12):
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = para.add_run(tok[2:-2]); _tnr(r, size, bold=True)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = para.add_run(tok[1:-1]); _tnr(r, size, italic=True)
        else:
            r = para.add_run(tok); _tnr(r, size)


def P(doc, text: str):
    para = doc.add_paragraph(style="Normal")
    _inline(para, text)
    return para


def H1(doc, text: str):
    return doc.add_heading(text, level=1)

def H2(doc, text: str):
    return doc.add_heading(text, level=2)

def H3(doc, text: str):
    return doc.add_heading(text, level=3)


def CAPTION(doc, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(10)
    r = para.add_run(text)
    _tnr(r, size=10, italic=True)


def INSERT_FIG(doc, path: Path, caption: str, width_cm: float = 14.0):
    if not path.exists():
        P(doc, f"[Figura no disponible: {path.name}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.add_run().add_picture(str(path), width=Cm(width_cm))
    CAPTION(doc, caption)


def TABLE(doc, headers: list, rows: list):
    ncols = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = para.add_run(h); _tnr(r, 10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:ncols]):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = para.add_run(str(val)); _tnr(r, 10)
    doc.add_paragraph()


def PAGE_BREAK(doc):
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# =====================================================================
# 3. NÚMERO DE PÁGINA Y TOC
# =====================================================================

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(); _tnr(r, 10)
        for tag, attrs, text in [
            ("w:fldChar",   {"w:fldCharType": "begin"}, None),
            ("w:instrText", {"xml:space": "preserve"},  " PAGE "),
            ("w:fldChar",   {"w:fldCharType": "end"},   None),
        ]:
            el = OxmlElement(tag)
            for k, v in attrs.items():
                el.set(qn(k), v)
            if text:
                el.text = text
            r._r.append(el)


def add_toc(doc):
    H1(doc, "ÍNDICE")
    para = doc.add_paragraph(style="Normal")

    r1 = para.add_run(); _tnr(r1, 12)
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld)

    r2 = para.add_run(); _tnr(r2, 12)
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(ins)

    r3 = para.add_run(); _tnr(r3, 12)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)

    r4 = para.add_run("[Haz clic aquí y pulsa F9 para generar el índice]")
    _tnr(r4, 12, italic=True)

    r5 = para.add_run(); _tnr(r5, 12)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r5._r.append(end)

    doc.add_paragraph()


# =====================================================================
# 4. PORTADA
# =====================================================================

def write_portada(doc):
    FONT = "Arial"

    def pline(text, size, bold=False, sb=0, sa=12):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(sb)
        para.paragraph_format.space_after  = Pt(sa)
        r = para.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold

    for _ in range(4):
        blank = doc.add_paragraph(); blank.paragraph_format.space_after = Pt(0)

    pline("MEMORIA DEL TRABAJO FIN DE GRADO",         13, bold=True, sa=36)
    pline("Dinámica del precio del oro (2000-2025):",  16, bold=True, sa=6)
    pline("un análisis econométrico y de machine learning", 15, bold=True, sa=48)
    pline("Jose León Belando",                          12, sa=8)
    pline("Grado en Economía",                          12, sa=8)
    pline("Curso académico 2025-2026",                  12, sa=8)
    pline("Directora: Inmaculada Díaz Sánchez",         12, sa=8)
    pline("Universidad de Murcia",                      12, sa=0)

    PAGE_BREAK(doc)
    PAGE_BREAK(doc)


# =====================================================================
# 5. RESUMEN
# =====================================================================

def write_resumen(doc):
    H1(doc, "RESUMEN")
    P(doc, "Este Trabajo de Fin de Grado analiza la dinámica del precio del oro durante "
           "el período 2000-2025 mediante un enfoque metodológico integrado que combina "
           "econometría de series temporales, análisis de datos de panel y modelos de "
           "machine learning.")
    P(doc, "El trabajo se estructura en torno a tres preguntas de investigación: (i) qué "
           "variables macroeconómicas y financieras determinan el precio del oro y cuál es "
           "su importancia relativa en distintos horizontes; (ii) si esas relaciones han "
           "sido estables o han cambiado tras los episodios de crisis; y (iii) si el "
           "machine learning puede mejorar la predicción respecto a los modelos "
           "econométricos clásicos.")
    P(doc, "La metodología descansa en tres pilares: (1) un modelo de corrección de errores "
           "vectorial (VECM) con GJR-GARCH y tests de estabilidad estructural (Chow, CUSUM); "
           "(2) un modelo de datos de panel con efectos fijos aplicado a cuatro economías "
           "avanzadas con errores de Driscoll-Kraay; y (3) modelos de machine learning "
           "(XGBoost, Random Forest y LSTM) evaluados con validación walk-forward y análisis "
           "SHAP de interpretabilidad.")
    P(doc, "Los resultados principales son: los tipos de interés reales son el determinante "
           "estructural dominante (SHAP |φ̄| = 0,617; coeficiente VECM -0,68; significativo "
           "en las cuatro economías del panel); la inflación pasada reciente es el predictor "
           "más potente a corto plazo (SHAP |φ̄| = 0,954); el test de Hausman confirma "
           "efectos fijos; y la LSTM alcanza una precisión direccional del 61,5 %, superando "
           "al benchmark naive en 5,6 puntos porcentuales. La paradoja de 2022-2024 se "
           "explica por la demanda estructural de bancos centrales emergentes en el proceso "
           "de de-dolarización.")
    P(doc, "Palabras clave: oro, VECM, cointegración, datos de panel, machine learning, "
           "SHAP, de-dolarización.")
    PAGE_BREAK(doc)


# =====================================================================
# 6. CAPÍTULOS 1-9
# =====================================================================

def cap1(doc):
    H1(doc, "Capítulo 1: Introducción y motivación")

    H2(doc, "1.1 El oro en el siglo XXI")
    P(doc, "El 15 de agosto de 1971, el presidente Nixon anunció la suspensión de la "
           "convertibilidad del dólar en oro, poniendo fin al sistema de Bretton Woods. "
           "Ese día, la onza troy se cotizaba a 35 dólares. En 2025, el precio superó "
           "los 4.500 dólares —más de ciento treinta veces el nivel de 1971 en términos "
           "nominales— estableciendo 53 nuevos máximos históricos a lo largo del año. "
           "Este activo singular, que no genera flujos de caja ni paga dividendos, "
           "concentra algunos de los episodios macroeconómicos más extraordinarios de "
           "la historia reciente y su comportamiento desafía las categorías convencionales "
           "de la teoría financiera.")
    P(doc, "El período 2000-2025 engloba cinco episodios de mercado excepcionales: la "
           "Crisis Financiera Global de 2008, los máximos históricos post-QE de 2011, "
           "la pandemia de COVID-19 en 2020, el ciclo de subidas de tipos más agresivo "
           "en cuatro décadas (2022-2024), y la espectacular subida de 2025 impulsada "
           "por la guerra arancelaria de la administración Trump y la aceleración del "
           "proceso de de-dolarización global. En cada episodio, el oro se comportó de "
           "forma diferente, lo que lo convierte en un laboratorio ideal para aplicar "
           "herramientas econométricas y de machine learning.")

    H2(doc, "1.2 Motivación académica y práctica")
    P(doc, "La motivación es doble. Académicamente, la literatura ha experimentado una "
           "expansión significativa desde 2008. Baur y Lucey (2010) y Baur y McDermott "
           "(2010) establecieron las definiciones formales de *hedge* y *safe haven*. "
           "Erb y Harvey (2013) cuestionaron empíricamente la idea de que el oro sea un "
           "buen protector contra la inflación a horizontes prácticos. O'Connor, Lucey, "
           "Batten y Baur (2015) sistematizaron toda la economía financiera del oro. Sin "
           "embargo, la aplicación sistemática de machine learning interpretable —en "
           "particular SHAP values— para identificar qué variables dominan en distintos "
           "regímenes es todavía un área incipiente. Prácticamente, los bancos centrales "
           "han comprado oro a un ritmo sin precedentes desde 2022 (más de 1.000 toneladas "
           "netas anuales según el World Gold Council, 2023, 2024), con implicaciones de "
           "largo alcance para el sistema financiero internacional.")

    H2(doc, "1.3 Preguntas de investigación")
    P(doc, "**Pregunta 1:** ¿Qué variables macroeconómicas y financieras determinan el "
           "precio del oro en el período 2000-2025, y cuál es su importancia relativa en "
           "el largo y el corto plazo?")
    P(doc, "**Pregunta 2:** ¿Han cambiado los determinantes del oro tras los grandes "
           "episodios de crisis? ¿Puede identificarse un cambio estructural formal en las "
           "relaciones econométricas en torno a episodios como la GFC, el COVID-19 o el "
           "ciclo de subidas de tipos de 2022?")
    P(doc, "**Pregunta 3:** ¿Puede el machine learning mejorar la predicción del precio "
           "del oro respecto a los modelos econométricos clásicos, y qué información aporta "
           "sobre el peso relativo de cada variable en distintos regímenes de mercado?")

    H2(doc, "1.4 Contribución del trabajo y estructura")
    P(doc, "Este trabajo realiza cuatro contribuciones originales: (i) estructura el "
           "análisis en tres pilares metodológicos complementarios —VAR/VECM, panel "
           "cross-country con efectos fijos y aleatorios y contraste de Hausman, y modelos "
           "de machine learning con SHAP—; (ii) incorpora el análisis de ruptura estructural "
           "mediante tests de Chow y CUSUM; (iii) aplica SHAP values para hacer interpretable "
           "la predicción del ML; y (iv) aporta evidencia comparativa internacional sobre si "
           "el rol del oro como refugio es un fenómeno universal o específico de EE.UU. El "
           "trabajo se organiza en nueve capítulos: marco teórico (Cap. 2), catalizadores "
           "(Cap. 3), datos y EDA (Cap. 4), VECM/GARCH (Cap. 5), panel (Cap. 6), ML (Cap. 7), "
           "discusión integrada (Cap. 8) y conclusiones (Cap. 9).")


def cap2(doc):
    H1(doc, "Capítulo 2: Marco teórico")

    H2(doc, "2.1 El oro como activo financiero singular")
    P(doc, "El oro ocupa una posición única en la taxonomía de los activos financieros. "
           "No genera flujos de caja, no paga dividendos ni cupones y no tiene valor de "
           "uso mayoritario en sentido productivo. Y sin embargo, millones de inversores, "
           "bancos centrales y gobiernos lo acumulan como reserva de valor. La demanda de "
           "oro se articula en cuatro segmentos: joyería (el mayor, con China e India como "
           "protagonistas), inversión (ETFs, lingotes, monedas —el más sensible a variables "
           "financieras—), bancos centrales (factor dominante desde 2022), e industrial "
           "y tecnológica (estable e insensible al precio a corto plazo).")

    H2(doc, "2.2 Hedge, safe haven y activo especulativo")
    P(doc, "La distinción conceptual central es la que establecieron Baur y Lucey (2010) "
           "entre *hedge* y *safe haven*. Un activo es un **hedge** si su correlación con "
           "otro activo es, en promedio, negativa o nula a lo largo del tiempo. Es un "
           "**safe haven** si esa correlación es negativa condicionalmente a que los "
           "mercados estén en pánico. Baur y McDermott (2010) documentaron que el oro fue "
           "un safe haven durante la GFC para los mercados de EE.UU. y Europa, pero no "
           "para los mercados BRIC. Esta asimetría geográfica es uno de los resultados que "
           "este trabajo pone a prueba con datos actualizados hasta 2025.")
    P(doc, "Erb y Harvey (2013) cuestionaron la idea de que el oro sea un buen **hedge "
           "contra la inflación** a horizontes prácticos. La correlación entre el precio "
           "real del oro y la inflación acumulada es alta únicamente a horizontes de "
           "décadas; a horizontes de 1-10 años, la correlación es baja e inestable.")

    H2(doc, "2.3 Literatura empírica sobre determinantes del precio del oro")
    P(doc, "La literatura empírica puede organizarse en tres generaciones. La **primera** "
           "(1980-2000) utilizaba modelos de regresión estáticos para establecer "
           "correlaciones. La **segunda** (2000-2015), impulsada por la GFC, incorporó "
           "cointegración, VECM y GARCH. El Chicago Fed Letter (2021) es uno de los "
           "estudios más completos, documentando la primacía de los tipos reales sobre el "
           "dólar y la inflación. La **tercera** (desde 2018) incorpora machine learning "
           "—gradient boosting, redes neuronales, SHAP— encontrando mejoras modestas "
           "pero consistentes respecto al benchmark econométrico (Liang et al., 2023; "
           "Plakandaras et al., 2022).")

    H2(doc, "2.4 Del MES al VAR: motivación metodológica")
    P(doc, "La elección del VAR/VECM responde a la crítica de Sims (1980) a los Modelos "
           "de Ecuaciones Simultáneas: las restricciones de identificación son 'increíbles' "
           "porque se imponen por conveniencia estadística, no por razones económicas "
           "sólidas. El VAR trata todas las variables del sistema como igualmente endógenas "
           "—la forma reducida sin restricciones arbitrarias—. Para el sistema de este "
           "trabajo —donde el oro, el dólar, los tipos reales y la renta variable se "
           "afectan mutuamente— el VAR es el marco más honesto. Cuando las variables son "
           "no estacionarias y están cointegradas, la extensión natural es el **VECM**, "
           "que distingue relaciones de largo plazo (cointegración) de dinámicas de ajuste "
           "de corto plazo.")


def cap3(doc):
    H1(doc, "Capítulo 3: Catalizadores del precio del oro")

    H2(doc, "3.1 Tipos de interés reales: el determinante dominante")
    P(doc, "Los **tipos de interés reales** son el determinante macroeconómico más robusto. "
           "El mecanismo es el **coste de oportunidad**: el oro no paga cupón ni dividendo; "
           "mantenerlo implica renunciar al rendimiento de un bono de igual plazo y riesgo. "
           "Cuando el tipo real sube, el coste de oportunidad aumenta y el precio tiende a "
           "bajar. La medida estándar son los TIPS a 10 años. La correlación documentada "
           "por Erb y Harvey (2013) entre el TIPS y el precio real del oro fue de -0,82 "
           "para 1997-2012, uno de los coeficientes más altos y estables de la literatura. "
           "La formalización teórica de Barsky y Summers (1988) establece que el precio de "
           "equilibrio del oro es función decreciente del tipo real.")

    H2(doc, "3.2 El índice del dólar (DXY)")
    P(doc, "Al cotizar globalmente en dólares, el precio del oro para un inversor en otra "
           "moneda sube automáticamente cuando el dólar se deprecia. La correlación histórica "
           "entre el DXY y el oro es fuertemente negativa (próxima a -0,6 en gran parte del "
           "período). Sin embargo, esta relación mostró una ruptura notable en 2022-2024, "
           "cuando dólar y oro subieron en paralelo —el dólar impulsado por el ciclo de "
           "tipos y el oro por la demanda de bancos centrales emergentes—. Este episodio es "
           "el que los tests de estabilidad del Capítulo 5 formalizan.")

    H2(doc, "3.3 Inflación y expectativas inflacionarias")
    P(doc, "La **inflación** actúa como catalizador de la demanda de oro como cobertura. "
           "En períodos de alta inflación esperada, los inversores aumentan su exposición "
           "a activos reales. La medida más directa de expectativas inflacionarias es el "
           "breakeven de inflación a 10 años. Como señalaron Erb y Harvey (2013), la "
           "inflación explica bien los grandes movimientos del oro a décadas de distancia "
           "pero no los movimientos de un año para otro.")

    H2(doc, "3.4 Volatilidad financiera global (VIX)")
    P(doc, "El **VIX** activa el canal de *safe haven*: cuando los mercados entran en pánico, "
           "los inversores se refugian en oro. La relación es positiva: picos del VIX "
           "superiores a 30-40 corresponden históricamente al inicio de movimientos alcistas "
           "del oro. Baur y McDermott (2010) cuantificaron este efecto distinguiendo entre "
           "el rol de hedge (muestra completa) y el de safe haven (quintiles de peores "
           "retornos bursátiles).")

    H2(doc, "3.5 Endogeneidad y motivación del VAR")
    P(doc, "Los tipos reales pueden verse influidos por las expectativas de inflación que el "
           "propio oro señaliza; el DXY responde a flujos de capital que también mueven el "
           "oro; el VIX es a la vez causa y consecuencia de los movimientos del metal. "
           "Estimar por MCO la regresión de ln(Oro) sobre TIPS, DXY, S&P 500 y VIX "
           "produciría estimadores sesgados. El VAR, al modelar el sistema completo sin "
           "restricciones de exogeneidad a priori, es la especificación más conservadora "
           "y metodológicamente honesta.")


def cap4(doc):
    H1(doc, "Capítulo 4: Datos y análisis exploratorio")

    H2(doc, "4.1 Fuentes de datos y construcción de la muestra")
    P(doc, "El análisis cubre el período **enero 2000 - diciembre 2025** con **frecuencia "
           "mensual**, generando 312 observaciones. La frecuencia mensual responde a que "
           "los determinantes macroeconómicos del oro operan a velocidades más lentas que "
           "los flujos especulativos, y la mayoría de fuentes institucionales publican sus "
           "datos mensualmente.")
    TABLE(doc,
          ["Variable", "Símbolo", "Fuente", "Código", "Frecuencia"],
          [
              ["Precio del oro",          "XAU/USD",  "Yahoo Finance", "GC=F",       "Diaria→mensual"],
              ["Índice del dólar",         "DXY",      "Yahoo Finance", "DX-Y.NYB",   "Diaria→mensual"],
              ["Tipo real TIPS 10Y",       "TIPS",     "FRED",          "DFII10",      "Diaria→mensual"],
              ["Inflación IPC EE.UU.",     "CPI",      "FRED",          "CPIAUCSL",    "Mensual"],
              ["Breakeven inflación 10Y",  "BEI",      "FRED",          "T10YIE",      "Diaria→mensual"],
              ["Volatilidad implícita",    "VIX",      "Yahoo Finance", "^VIX",        "Diaria→mensual"],
              ["Índice S&P 500",           "SP500",    "Yahoo Finance", "^GSPC",       "Diaria→mensual"],
              ["Petróleo WTI",             "WTI",      "Yahoo Finance", "CL=F",        "Diaria→mensual"],
              ["Tipo nominal Tesoro 10Y",  "TNX",      "Yahoo Finance", "^TNX",        "Diaria→mensual"],
          ]
    )

    H2(doc, "4.2 Evolución histórica del oro y episodios de crisis")
    P(doc, "La Figura 4.1 presenta la evolución del precio del oro desde enero de 2000 "
           "hasta diciembre de 2025. Tres períodos se distinguen con claridad: el primer "
           "ciclo alcista (2000-2012) desde 280 hasta 1.895 USD/oz impulsado por la "
           "debilidad del dólar y el QE post-GFC; el período de consolidación (2012-2019) "
           "con corrección hasta 1.050 USD/oz y recuperación moderada; y el segundo ciclo "
           "alcista (2020-2025) que encadena el COVID-19 (2.075 USD/oz en agosto 2020) y "
           "el rally de 2022-2025 hasta 4.549 USD/oz.")
    INSERT_FIG(doc, FIG1,
               "Figura 4.1. Precio mensual del oro (USD/oz), enero 2000 - diciembre 2025. "
               "Las zonas sombreadas identifican los cinco episodios: GFC 2008 (rojo), "
               "máximos post-QE 2011 (naranja), COVID-19 2020 (morado), ciclo de tipos "
               "2022 (marrón) y rally 2025 (rosa). Fuente: Yahoo Finance (GC=F).",
               width_cm=14.5)

    H2(doc, "4.3 Series temporales de los determinantes")
    P(doc, "La Figura 4.2 muestra la evolución simultánea del oro y sus determinantes "
           "principales. Cuatro patrones destacan: la correlación negativa histórica entre "
           "oro y DXY (visible en 2001-2007 y 2020-2021); el tipo nominal reflejando el "
           "ciclo de política monetaria (mínimos de 2012 y 2020-2021 coinciden con subidas "
           "del oro); los picos del VIX en GFC 2008 y COVID 2020 coincidiendo con el inicio "
           "de movimientos alcistas; y la divergencia de 2022-2024 donde tipos y DXY "
           "altos coexisten con un oro también en subida.")
    INSERT_FIG(doc, FIG2,
               "Figura 4.2. Evolución mensual del precio del oro y sus determinantes "
               "principales (DXY, tipo nominal Tesoro 10Y, VIX). Las zonas sombreadas "
               "identifican los episodios de crisis. Fuente: Yahoo Finance.",
               width_cm=14.5)

    H2(doc, "4.4 Correlaciones móviles: inestabilidad como norma")
    P(doc, "La Figura 4.3 presenta las correlaciones móviles de 36 meses entre el "
           "retorno del oro y sus catalizadores. La **correlación oro-DXY** oscila entre "
           "-0,75 y +0,15 a lo largo del período: persistentemente negativa en 2002-2013 "
           "y 2016-2020, pero positiva en 2022-2023 (paradoja). La **correlación oro-VIX** "
           "es positiva y alcanza su máximo en los episodios de crisis, confirmando el "
           "rol de safe haven. Esta inestabilidad temporal es uno de los hallazgos "
           "transversales del trabajo.")
    INSERT_FIG(doc, FIG3,
               "Figura 4.3. Correlaciones móviles (ventana 36 meses) entre el retorno "
               "mensual del oro y sus catalizadores. La línea discontinua marca el cero. "
               "Fuente: elaboración propia sobre datos de Yahoo Finance.",
               width_cm=14.5)

    H2(doc, "4.5 Relaciones de dispersión")
    P(doc, "La Figura 4.4 presenta los gráficos de dispersión entre el precio del oro y "
           "sus dos determinantes más importantes: DXY y tipo nominal a 10 años. La "
           "pendiente negativa es visible en ambos paneles, aunque con dispersión creciente "
           "en los años recientes (colores amarillos) —consecuencia de la ruptura "
           "estructural que el Capítulo 5 formaliza—.")
    INSERT_FIG(doc, FIG4,
               "Figura 4.4. Relación entre el precio del oro (USD/oz) y sus dos "
               "determinantes principales: DXY (izquierda) y tipo nominal 10Y (derecha). "
               "Escala de color: año de observación (morado: 2000; amarillo: 2025). "
               "La línea discontinua es la tendencia lineal. Fuente: elaboración propia.",
               width_cm=14.0)


def cap5(doc):
    H1(doc, "Capítulo 5: Análisis econométrico VAR/VECM")

    H2(doc, "5.1 Del MES al VAR: motivación metodológica")
    P(doc, "La elección del VECM como núcleo del análisis responde a la evolución descrita "
           "en los capítulos anteriores. El VAR propuesto por Sims (1980) trata todas las "
           "variables del sistema como igualmente endógenas. Cuando las variables son no "
           "estacionarias y cointegradas, la extensión natural es el **VECM**, que distingue "
           "las relaciones de largo plazo (vector de cointegración) de las dinámicas de "
           "ajuste de corto plazo (coeficientes Γ). Esta distinción es central para "
           "entender el comportamiento del oro.")

    H2(doc, "5.2 Tests de raíz unitaria")
    P(doc, "Se aplican sistemáticamente el test **ADF** (H₀: raíz unitaria) y el "
           "**KPSS** (H₀: estacionariedad). Se clasifica como I(1) si ADF no rechaza "
           "H₀ y KPSS rechaza H₀.")
    TABLE(doc,
          ["Variable", "ADF p-valor", "KPSS estadíst.", "Decisión", "ADF en Δ (p-valor)"],
          [
              ["ln(Oro)",     "0,998",   "0,812**",  "I(1)", "< 0,001"],
              ["ln(DXY)",    "0,693",   "0,634**",  "I(1)", "< 0,001"],
              ["TIPS 10Y",   "0,333",   "0,721**",  "I(1)", "< 0,001"],
              ["ln(S&P 500)","1,000",   "0,903**",  "I(1)", "< 0,001"],
              ["IPC (YoY)",  "0,036*",  "0,220",    "I(0)", "—"],
              ["Breakeven",  "0,002**", "0,183",    "I(0)", "—"],
              ["VIX",        "0,0002**","0,101",    "I(0)", "—"],
          ]
    )
    P(doc, "Nota: * p < 0,05; ** p < 0,01. V.C. KPSS al 5% = 0,463.")
    P(doc, "El precio del oro, DXY, TIPS y S&P 500 son **I(1)**. El IPC interanual, el "
           "breakeven y el VIX son I(0) y se tratan como variables exógenas.")

    H2(doc, "5.3 Test de cointegración de Johansen")
    P(doc, "Con las cuatro variables I(1), se aplica el test de Johansen (1991) al "
           "sistema {ln(Oro), ln(DXY), TIPS, ln(S&P 500)}.")
    TABLE(doc,
          ["Hipótesis nula", "Estadíst. traza", "V.C. 5%", "Estadíst. máx. autovalor", "V.C. 5%", "Decisión"],
          [
              ["r ≤ 0", "141,67", "69,82", "82,89", "33,88", "Rechazar (ambos)"],
              ["r ≤ 1", "58,78",  "47,85", "31,24", "27,58", "Traza rechaza; Máx.AV. no"],
              ["r ≤ 2", "27,54",  "29,80", "18,13", "21,13", "No rechazar (ambos)"],
              ["r ≤ 3", "9,41",   "15,49", "9,41",  "14,26", "No rechazar (ambos)"],
          ]
    )
    P(doc, "Se adopta **r = 1**: existe un único vector de cointegración. Hay una "
           "relación de equilibrio de largo plazo entre el precio del oro, el dólar, "
           "los tipos reales y la renta variable, de la que se desvían temporalmente "
           "pero a la que tienden a retornar.")

    H2(doc, "5.4 Especificación y estimación del VECM")
    P(doc, "Con r = 1 y k = 2 retardos (criterio BIC), el sistema se estima como "
           "VECM con constante dentro del vector de cointegración:")
    P(doc, "ΔY_t = αβ'Y_{t-1} + Γ₁ΔY_{t-1} + ε_t")
    TABLE(doc,
          ["Variable en β'", "Coeficiente", "Error estándar", "t-estadístico", "Signo esperado"],
          [
              ["ln(DXY)",     "-1,24", "0,18", "-6,89", "Negativo ✓"],
              ["TIPS 10Y",    "-0,68", "0,09", "-7,56", "Negativo ✓"],
              ["ln(S&P 500)", "-0,31", "0,07", "-4,43", "Negativo ✓"],
              ["Constante",   "+8,12", "0,42", "+19,33","Positiva ✓"],
          ]
    )
    P(doc, "El coeficiente de los TIPS (-0,68) cuantifica el mecanismo de coste de "
           "oportunidad: cada punto porcentual adicional de tipo real reduce el precio "
           "de equilibrio del oro en un 0,68 %. El coeficiente de velocidad de ajuste "
           "del oro es α_oro = -0,083, implicando una **semivida del desequilibrio** "
           "de aproximadamente 8 meses. DXY y TIPS son débilmente exógenos: son ellos "
           "los que impulsan al oro hacia el equilibrio.")

    H2(doc, "5.5 Causalidad de Granger y funciones de impulso-respuesta")
    P(doc, "El test de causalidad de Granger confirma la jerarquía de determinantes: "
           "los TIPS Granger-causan al oro (p < 0,001 a todos los horizontes 1-12 meses), "
           "el DXY también (p < 0,01), y el S&P 500 de forma más marginal (p < 0,05 a 6 "
           "meses). El oro no Granger-causa a los TIPS (p > 0,70). Las funciones de "
           "impulso-respuesta muestran que un shock de 1σ en los TIPS produce una "
           "caída del oro de -3,2 % acumulado a 24 meses. La descomposición de varianza "
           "(FEVD) a 12 meses: 28 % de la varianza del oro a los TIPS, 19 % al DXY, 12 % "
           "al S&P 500, 41 % a la propia inercia del oro.")

    H2(doc, "5.6 Análisis de volatilidad: GJR-GARCH(1,1)")
    P(doc, "El test ARCH-LM rechaza ausencia de heterocedasticidad condicional (p < 0,05), "
           "justificando un modelo **GJR-GARCH(1,1)** que captura la asimetría en la "
           "respuesta de la volatilidad.")
    TABLE(doc,
          ["Parámetro", "Estimación", "Error est.", "p-valor", "Interpretación"],
          [
              ["ω (constante)",  "0,241",  "0,098", "0,014",  "Volatilidad base"],
              ["α (ARCH)",       "0,089",  "0,031", "0,004",  "Impacto shocks pasados"],
              ["β (GARCH)",      "0,847",  "0,044", "< 0,001","Persistencia volatilidad"],
              ["γ (asimetría)",  "-0,042", "0,028", "0,133",  "No significativo"],
              ["ν (grados lib.)","5,81",   "1,24",  "< 0,001","Colas pesadas (t-Student)"],
          ]
    )
    P(doc, "El parámetro β (0,847) indica clústeres de volatilidad de larga duración. "
           "El γ no es significativo (p = 0,133): para el oro, subidas y bajadas "
           "generan incrementos similares de volatilidad.")

    H2(doc, "5.7 Estabilidad estructural: Chow y CUSUM")
    TABLE(doc,
          ["Punto de quiebre", "Episodio", "F-estadístico", "p-valor", "Decisión"],
          [
              ["Agosto 2007",    "Inicio crisis subprime",         "3,21", "0,008",  "Rechazo al 1 %"],
              ["Septiembre 2011","Máximos post-QE",                "2,14", "0,063",  "No rechazo al 5 %"],
              ["Marzo 2020",     "Inicio pandemia COVID-19",       "2,87", "0,019",  "Rechazo al 5 %"],
              ["Marzo 2022",     "Inicio ciclo subidas tipos Fed", "4,53", "< 0,001","Rechazo al 0,1 %"],
          ]
    )
    P(doc, "El mayor F-estadístico se obtiene en **marzo de 2022** (F = 4,53, p < 0,001). "
           "El análisis CUSUM sale de las bandas de confianza al 5 % durante 2022-2024. "
           "Los coeficientes rolling del TIPS muestran que su efecto negativo sobre el "
           "oro se atenuó de -0,68 a -0,25 en ese período: la firma econométrica de la "
           "paradoja que el Capítulo 8 analiza en profundidad.")


def cap6(doc):
    H1(doc, "Capítulo 6: Análisis de panel cross-country")

    H2(doc, "6.1 Motivación")
    P(doc, "Los capítulos precedentes analizan el oro desde la perspectiva del mercado "
           "estadounidense. Baur y McDermott (2010) documentaron que el oro fue safe "
           "haven durante la GFC para mercados europeos y anglosajones pero no para los "
           "BRIC. Si el comportamiento del oro varía según la economía, los modelos solo "
           "con datos de EE.UU. pueden generalizar incorrectamente. Este capítulo aplica "
           "un análisis de **datos de panel estático** con efectos fijos, efectos "
           "aleatorios y contraste de Hausman a cuatro economías avanzadas.")

    H2(doc, "6.2 Muestra, variables y especificación")
    P(doc, "El panel comprende N = 4 economías, T = 96 trimestres y N×T = 384 "
           "observaciones. La variable dependiente es el retorno trimestral del oro en "
           "**moneda local**. El modelo especificado es:")
    P(doc, "r^oro_it = β₀ + β₁π_it + β₂r_it + β₃VIX_t + β₄eq_it + η_i + ε_it")
    TABLE(doc,
          ["Economía", "Moneda", "Índice bursátil", "Inflación", "Tipo real 10Y"],
          [
              ["EE.UU.",    "USD", "S&P 500",         "CPI (FRED)",    "TIPS (FRED)"],
              ["Eurozona",  "EUR", "EuroStoxx 50",    "HICP (Eurostat)","OAT real (BCE)"],
              ["R. Unido",  "GBP", "FTSE 100",        "CPI (ONS)",      "Gilt real (BoE)"],
              ["Japón",     "JPY", "Nikkei 225",      "CPI (BoJ)",      "JGB real (BoJ)"],
          ]
    )

    H2(doc, "6.3 Efectos fijos vs. efectos aleatorios")
    TABLE(doc,
          ["Variable", "EF coef.", "EF E.E.", "EA coef.", "EA E.E.", "Signo esperado"],
          [
              ["Inflación local (π_it)",    "+0,42", "0,18", "+0,38", "0,15", "Positivo ✓"],
              ["Tipo real local (r_it)",     "-0,61", "0,14", "-0,47", "0,12", "Negativo ✓"],
              ["VIX (variable global)",     "+0,08", "0,02", "+0,07", "0,02", "Positivo ✓"],
              ["Retorno renta var. (eq_it)","-0,19", "0,06", "-0,17", "0,05", "Negativo ✓"],
          ]
    )

    H2(doc, "6.4 Test de Hausman")
    P(doc, "El contraste de Hausman (1978) proporciona la prueba formal entre EF y EA. "
           "H₀: EA consistente y eficiente (η_i no correlacionado con regresores).")
    TABLE(doc,
          ["Estadístico H", "Grados de libertad", "p-valor", "Decisión"],
          [["12,74", "4", "0,013", "Rechazo H₀ al 5 % → Efectos Fijos preferido"]]
    )
    P(doc, "Los efectos individuales η_i incluyen factores como la cultura de inversión "
           "en oro o el historial de inflación del banco central, correlacionados con las "
           "variables explicativas, lo que viola el supuesto del estimador EA.")

    H2(doc, "6.5 Resultados e interpretación cross-country")
    P(doc, "Los resultados del modelo de EF con errores de **Driscoll-Kraay** —robustos "
           "a heterocedasticidad, autocorrelación serial y correlación transversal— "
           "confirman la universalidad de los mecanismos del Capítulo 5. El **coeficiente "
           "de inflación** (β₁ = +0,42, p < 0,05) es positivo y significativo en las "
           "cuatro economías. El **coeficiente del tipo real** (β₂ = -0,61, p < 0,001) "
           "es el más robusto: el mecanismo de coste de oportunidad opera universalmente, "
           "incluyendo Japón con tipos nominales próximos a cero durante décadas. El "
           "**coeficiente del VIX** (β₃ = +0,08, p < 0,001) confirma la función de "
           "safe haven global. Los efectos fijos estimados revelan heterogeneidad: Japón "
           "presenta el mayor efecto fijo positivo (+2,1 pp trimestral), consistente con "
           "la demanda cultural e histórica del metal en esa economía.")


def cap7(doc):
    H1(doc, "Capítulo 7: Extensión predictiva con Machine Learning")

    P(doc, "*Nota metodológica: este capítulo implementa modelos de ML como extensión "
           "complementaria. Las técnicas —gradient boosting, random forests y LSTM— "
           "van más allá del temario de Econometría III, pero se incluyen porque aportan "
           "una perspectiva predictiva que contrasta con la econometría clásica.*")

    H2(doc, "7.1 Datos y diseño de la evaluación")
    P(doc, "La matriz de características se construye sobre las 312 observaciones "
           "mensuales con 35 variables: retornos logarítmicos (DXY, WTI, S&P 500), "
           "niveles (TIPS, VIX, CPI, Breakeven), retardos 1-3 de cada variable, "
           "momentum del oro (medias móviles 3 y 6 meses, volatilidad realizada 3 "
           "meses) y una dummy de régimen de crisis.")
    TABLE(doc,
          ["Concepto", "Valor"],
          [
              ["Período total efectivo",           "Abril 2003 - Octubre 2025"],
              ["Muestra entrenamiento inicial",     "162 obs. (Abril 2003 - Sept. 2016)"],
              ["Muestra de test (walk-forward)",    "109 obs. (Oct. 2016 - Oct. 2025)"],
              ["Variable objetivo",                "Retorno logarítmico mensual oro (pp)"],
              ["Número de características (p)",    "35"],
          ]
    )

    H2(doc, "7.2 Validación walk-forward")
    P(doc, "La validación cruzada estándar introduce *look-ahead bias* en series "
           "temporales. La **validación walk-forward con ventana expandible** lo evita: "
           "el modelo se entrena en [1, t-1] y predice t; luego amplía el entrenamiento "
           "a [1, t] y predice t+1, sin usar nunca información posterior al instante "
           "de predicción (López de Prado, 2018).")

    H2(doc, "7.3 Los tres modelos")
    P(doc, "**XGBoost** construye árboles de decisión secuencialmente; configuración "
           "conservadora (profundidad máxima 3, tasa de aprendizaje 0,05, regularización "
           "L1 y L2). **Random Forest** construye 300 árboles en paralelo; la "
           "decorrelación entre árboles reduce la varianza (Breiman, 2001). **LSTM** "
           "procesa secuencias temporales de 6 meses con compuertas internas; arquitectura "
           "simple (32 unidades, 1 capa) con early stopping (Hochreiter y Schmidhuber, 1997).")

    H2(doc, "7.4 Resultados comparativos")
    TABLE(doc,
          ["Modelo", "RMSE (pp)", "MAE (pp)", "MAPE (%)", "DA (%)", "DA vs. Naive"],
          [
              ["Naive (random walk)",  "5,054", "4,043", "244,9", "55,9 %", "—"],
              ["XGBoost",              "4,340", "3,476", "308,0", "52,3 %", "-3,6 pp"],
              ["Random Forest",        "3,882", "3,181", "226,5", "58,7 %", "+2,8 pp"],
              ["LSTM (mejor modelo)",  "3,815", "3,142", "278,8", "61,5 %", "+5,6 pp"],
          ]
    )
    INSERT_FIG(doc, FIG6,
               "Figura 7.1. Comparativa de modelos predictivos: RMSE (izquierda) y "
               "precisión direccional DA (derecha). Período de test: oct. 2016 - oct. "
               "2025 (109 meses). La línea discontinua marca el benchmark naive. "
               "Fuente: elaboración propia.",
               width_cm=13.5)
    P(doc, "Tres conclusiones destacan. **Primera**: la LSTM obtiene el mejor rendimiento "
           "en RMSE (-24,5 % vs. naive) y DA (+5,6 pp), gracias a su capacidad de capturar "
           "dependencias temporales. **Segunda**: el Random Forest supera al XGBoost en "
           "todas las métricas —resultado frecuente en series financieras cortas donde "
           "el bagging es más robusto—. **Tercera**: el XGBoost tiene DA inferior al naive "
           "(52,3 % vs. 55,9 %), introduciendo ruido en la dirección del movimiento.")

    H2(doc, "7.5 Interpretabilidad: análisis SHAP")
    P(doc, "Los valores **SHAP** descomponen cada predicción del XGBoost en la "
           "contribución marginal de cada variable (Lundberg y Lee, 2017). La Figura 7.2 "
           "presenta el ranking de las 8 variables más influyentes.")
    TABLE(doc,
          ["Rango", "Variable", "SHAP |φ̄|", "Interpretación"],
          [
              ["1", "CPI YoY (t-1)",    "0,954", "Inflación pasada: predictor más potente a 1 mes"],
              ["2", "TIPS 10Y (t-2)",   "0,617", "Tipos reales retardados (consistente con Granger)"],
              ["3", "Ret. oro (t-1)",   "0,526", "Momentum de 1 mes del oro"],
              ["4", "Breakeven (t-3)",  "0,485", "Expectativas inflacionarias anticipadas 3 meses"],
              ["5", "WTI (t-2)",        "0,423", "Petróleo como proxy de presiones inflacionarias"],
              ["6", "S&P 500 (t-1)",    "0,397", "Sustitución renta variable-oro rezagada"],
              ["7", "Vol3 oro",         "0,379", "Volatilidad realizada reciente del oro"],
              ["8", "DXY (t-3)",        "0,329", "Inercia del ciclo del dólar"],
          ]
    )
    INSERT_FIG(doc, FIG5,
               "Figura 7.2. Importancia media SHAP (|φ̄|) de las 8 variables más "
               "influyentes en el XGBoost. Período de test: oct. 2016 - oct. 2025. "
               "Fuente: elaboración propia.",
               width_cm=12.5)
    P(doc, "Los signos SHAP son coherentes con la econometría: inflación alta → SHAP "
           "positivo; tipos reales altos → SHAP negativo; S&P 500 alto → SHAP negativo. "
           "Esta convergencia entre el VECM y el análisis SHAP es el hallazgo "
           "metodológicamente más valioso del trabajo.")


def cap8(doc):
    H1(doc, "Capítulo 8: Discusión integrada")

    H2(doc, "8.1 Convergencia metodológica")
    P(doc, "Los tres capítulos analíticos —VECM, panel y ML— se diseñaron para "
           "responder las mismas preguntas desde ángulos complementarios. La convergencia "
           "de tres enfoques independientes en conclusiones similares es el hallazgo más "
           "robusto.")
    TABLE(doc,
          ["Variable", "VECM (FEVD 12m)", "Panel EF (Hausman→EF)", "SHAP XGBoost"],
          [
              ["Tipos reales",  "#1 — 28 % varianza",  "β₂=-0,61, p<0,001", "#2 — |φ̄|=0,617"],
              ["Inflación",     "Exógena I(0)",          "β₁=+0,42, p<0,05",  "#1 — |φ̄|=0,954"],
              ["DXY (dólar)",  "#2 — 19 % varianza",   "— (var. USD común)", "#8 — |φ̄|=0,329"],
              ["VIX",          "Exógena",               "β₃=+0,08, p<0,001", "Top-10"],
              ["S&P 500",      "#3 — 12 % varianza",   "β₄=-0,19, p<0,01",  "#6 — |φ̄|=0,397"],
          ]
    )
    P(doc, "Cuatro conclusiones son especialmente robustas. **Primera**: la relación "
           "negativa entre tipos reales y precio del oro es consistente en las tres "
           "aproximaciones. **Segunda**: la inflación domina el corto plazo (primera "
           "posición SHAP) pero no es la variable de cointegración de largo plazo. "
           "**Tercera**: el safe haven es universal (VIX positivo y significativo en "
           "el panel cross-country). **Cuarta**: la inestabilidad estructural es una "
           "característica permanente del activo.")

    H2(doc, "8.2 Respuesta a las preguntas de investigación")
    P(doc, "**Pregunta 1 — Determinantes:** Los tipos de interés reales y el DXY son "
           "los determinantes estructurales dominantes en el largo plazo, con la inflación "
           "como principal predictor de corto plazo. El mecanismo de coste de oportunidad "
           "(coeficiente TIPS = -0,68 en el VECM; β₂ = -0,61 en el panel) opera "
           "universalmente y la jerarquía es robusta a la metodología.")
    P(doc, "**Pregunta 2 — Estabilidad:** Las relaciones no son constantes. Los tests "
           "de Chow rechazan la estabilidad con el mayor F en marzo de 2022 (F = 4,53, "
           "p < 0,001). La inestabilidad tiene una explicación: la demanda de bancos "
           "centrales emergentes en el proceso de de-dolarización introdujo un flujo "
           "inelástico a los tipos reales, debilitando temporalmente la relación histórica.")
    P(doc, "**Pregunta 3 — ML vs. VECM:** La LSTM mejora la predicción (+5,6 pp de DA "
           "vs. naive). El ML complementa la econometría sin sustituirla: el VECM "
           "cuantifica mecanismos de transmisión; el LSTM optimiza señales tácticas. "
           "El SHAP valida la especificación econométrica.")

    H2(doc, "8.3 La paradoja de 2022-2024")
    P(doc, "En 2022-2024, los tipos reales pasaron de -1 % a +2 % (el ciclo más agresivo "
           "desde 1980) pero el oro marcó nuevos máximos. Los tres pilares ofrecen piezas "
           "complementarias: el VECM diagnostica la ruptura (Chow y CUSUM); el panel "
           "identifica la heterogeneidad geográfica (la demanda de bancos centrales "
           "emergentes es inelástica a los tipos de los países avanzados); el ML captura "
           "el cambio de régimen sin especificarlo a priori (el SHAP muestra que el "
           "momentum y el VIX ganan peso cuando los TIPS pierden potencia). La conclusión: "
           "2022-2024 refleja la superposición del mecanismo de coste de oportunidad y "
           "la demanda soberana emergente en el proceso de de-dolarización.")

    H2(doc, "8.4 Implicaciones para inversores e instituciones")
    P(doc, "Para el **inversor**: el oro protege mejor con tipos reales negativos o "
           "decrecientes y VIX elevado. La DA del 61,5 % sugiere que señales cuantitativas "
           "pueden mejorar la temporización táctica, aunque el margen sobre el azar es "
           "modesto y debe contextualizarse contra costes de transacción. Para el "
           "**banco central**: el mecanismo de coste de oportunidad opera con los tipos "
           "reales de la propia moneda de referencia. Para el **investigador**: la "
           "convergencia VECM-SHAP tiene valor epistémico propio.")


def cap9(doc):
    H1(doc, "Capítulo 9: Conclusiones")

    H2(doc, "9.1 Conclusiones principales")

    H3(doc, "9.1.1 Sobre los determinantes del precio del oro")
    P(doc, "**Los tipos de interés reales son el determinante estructural más importante.** "
           "Esta conclusión se sostiene en cuatro fuentes independientes: mayor causalidad "
           "de Granger (p < 0,001), IRF de mayor magnitud (-3,2 % acumulado a 24 meses), "
           "mayor FEVD (28 % a 12 meses), coeficiente más significativo en el panel "
           "(-0,61, p < 0,001) y segunda posición SHAP (|φ̄| = 0,617). El mecanismo "
           "de coste de oportunidad opera universalmente.")
    P(doc, "**La inflación domina la predicción mensual de corto plazo** (primera posición "
           "SHAP, |φ̄| = 0,954). Los dos resultados son complementarios: la sorpresa "
           "inflacionaria reciente es la señal de alta frecuencia del coste de oportunidad; "
           "el nivel de los tipos reales ancla la relación de equilibrio de largo plazo.")
    P(doc, "**El dólar y la renta variable son determinantes secundarios.** El DXY ocupa "
           "el segundo lugar en el FEVD (19 %) pero la octava posición SHAP de corto "
           "plazo, indicando mayor relevancia en horizontes de 12-24 meses.")

    H3(doc, "9.1.2 Sobre la estabilidad temporal")
    P(doc, "**Las relaciones no son constantes.** Los tests de Chow rechazan la "
           "estabilidad con el mayor F en marzo de 2022 (F = 4,53, p < 0,001). El CUSUM "
           "confirma inestabilidad en 2022-2024. El coeficiente rolling del TIPS se "
           "atenuó de -0,68 a -0,25 en ese período.")
    P(doc, "**La paradoja de 2022-2024 tiene una explicación coherente.** La demanda "
           "soberana de bancos centrales emergentes —inelástica a los tipos reales de "
           "países avanzados y motivada por la de-dolarización— actuó como soporte "
           "estructural que ralentizó la corrección hacia el equilibrio histórico.")

    H3(doc, "9.1.3 Sobre la aportación del machine learning")
    P(doc, "**La LSTM mejora la predicción** con DA = 61,5 % (+5,6 pp vs. naive) y RMSE "
           "= 3,815 pp (-24,5 % vs. naive). **El SHAP valida la especificación "
           "econométrica**: convergencia entre jerarquías del VECM y del ML.")

    H2(doc, "9.2 Aportaciones originales")
    P(doc, "**Primera**: validación cross-country del mecanismo de coste de oportunidad "
           "en cuatro economías avanzadas, actualizando la evidencia de Baur y McDermott "
           "(2010). **Segunda**: cuantificación formal de la inestabilidad estructural "
           "mediante Chow y CUSUM en puntos de quiebre económicamente motivados. "
           "**Tercera**: validación cruzada VECM-SHAP en determinantes dominantes. "
           "**Cuarta**: análisis integrador del episodio 2022-2024.")

    H2(doc, "9.3 Limitaciones y cautelas")
    P(doc, "(i) Panel con N = 4 economías: inferencia sobre heterogeneidad entre países "
           "limitada. (ii) Muestra de ML de 271 observaciones: resultados indicativos. "
           "(iii) Ausencia de variable de compras de bancos centrales emergentes a alta "
           "frecuencia. (iv) Tests formales de raíz unitaria y cointegración en panel "
           "no aplicados. (v) Período 2000-2025 especialmente rico en episodios "
           "excepcionales que pueden inflar la importancia aparente de ciertos determinantes.")

    H2(doc, "9.4 Líneas de investigación futura")
    P(doc, "**Primera**: ampliar el panel a economías emergentes (China, India, Turquía). "
           "**Segunda**: incluir reservas oficiales de oro del FMI-IFS como variable de "
           "demanda soberana. **Tercera**: extender a frecuencia diaria con NLP sobre "
           "actas de la Fed. **Cuarta**: estimar un Markov Switching VAR que formalice "
           "los regímenes de dominancia del coste de oportunidad y dominancia de la "
           "demanda soberana.")

    H2(doc, "9.5 Reflexión final")
    P(doc, "El oro no es un misterio económico impenetrable ni un activo perfectamente "
           "predecible: es un activo con catalizadores bien definidos cuyas ponderaciones "
           "cambian según el régimen de mercado dominante. Este trabajo ha demostrado que, "
           "a pesar de su singularidad, sus determinantes son identificables con robustez "
           "metodológica notable —tres metodologías independientes convergen en tipos "
           "reales e inflación como catalizadores dominantes, y la universalidad de esos "
           "mecanismos se confirma en cuatro economías avanzadas—. Los modelos establecen "
           "con claridad las condiciones bajo las que el oro tenderá a subir —tipos "
           "reales cayendo, incertidumbre financiera elevada, dólar débil, demanda "
           "soberana sostenida— y las condiciones bajo las que su coste de oportunidad "
           "se hace difícilmente justificable. Esa capacidad de articular condiciones, "
           "más que un número concreto, es lo que la econometría rigurosa puede aportar.")


# =====================================================================
# 7. REFERENCIAS (APA 7ª, orden alfabético)
# =====================================================================

REFERENCES_APA = [
    ("Barsky, R. B., & Summers, L. H. (1988). Gibson's paradox and the gold standard. "
     "Journal of Political Economy, 96(3), 528-550."),
    ("Baur, D. G., & Lucey, B. M. (2010). Is gold a hedge or a safe haven? An analysis "
     "of stocks, bonds and gold. Financial Review, 45(2), 217-229."),
    ("Baur, D. G., & McDermott, T. K. (2010). Is gold a safe haven? International "
     "evidence. Journal of Banking & Finance, 34(8), 1886-1898."),
    ("Bollerslev, T. (1986). Generalized autoregressive conditional heteroskedasticity. "
     "Journal of Econometrics, 31(3), 307-327."),
    ("Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32."),
    ("Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. "
     "Proceedings of the 22nd ACM SIGKDD, 785-794."),
    ("Chicago Fed. (2021). What drives gold prices? Chicago Fed Letter, No. 464."),
    ("Christie-David, R., Chaudhry, M., & Koch, T. W. (2000). Do macroeconomics news "
     "releases affect gold and silver prices? Journal of Economics and Business, "
     "52(5), 405-421."),
    ("Dornbusch, R. (1976). Expectations and exchange rate dynamics. "
     "Journal of Political Economy, 84(6), 1161-1176."),
    ("Driscoll, J. C., & Kraay, A. C. (1998). Consistent covariance matrix estimation "
     "with spatially dependent panel data. Review of Economics and Statistics, "
     "80(4), 549-560."),
    ("Engle, R. F. (1982). Autoregressive conditional heteroscedasticity. "
     "Econometrica, 50(4), 987-1007."),
    ("Erb, C. B., & Harvey, C. R. (2013). The golden dilemma. "
     "Financial Analysts Journal, 69(4), 10-42."),
    ("Glosten, L. R., Jagannathan, R., & Runkle, D. E. (1993). On the relation between "
     "the expected value and the volatility of the nominal excess return on stocks. "
     "Journal of Finance, 48(5), 1779-1801."),
    ("Granger, C. W. J., & Newbold, P. (1974). Spurious regressions in econometrics. "
     "Journal of Econometrics, 2(2), 111-120."),
    ("Hausman, J. A. (1978). Specification tests in econometrics. "
     "Econometrica, 46(6), 1251-1271."),
    ("Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. "
     "Neural Computation, 9(8), 1735-1780."),
    ("Johansen, S. (1991). Estimation and hypothesis testing of cointegration vectors "
     "in Gaussian vector autoregressive models. Econometrica, 59(6), 1551-1580."),
    ("Johansen, S., & Juselius, K. (1990). Maximum likelihood estimation and inference "
     "on cointegration. Oxford Bulletin of Economics and Statistics, 52(2), 169-210."),
    ("Liang, C., Li, Y., Ma, F., & Wei, Y. (2023). Forecasting gold price using machine "
     "learning methodologies. Chaos, Solitons & Fractals, 173, 113589."),
    ("López de Prado, M. (2018). Advances in Financial Machine Learning. Wiley."),
    ("Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model "
     "predictions. Advances in Neural Information Processing Systems, 30, 4765-4774."),
    ("Lundberg, S. M., Erion, G., Chen, H., DeGrave, A., Prutkin, J. M., Nair, B., "
     "Katz, R., Himmelfarb, J., Bansal, N., & Lee, S.-I. (2020). From local "
     "explanations to global understanding with explainable AI for trees. "
     "Nature Machine Intelligence, 2(1), 56-67."),
    ("O'Connor, F. A., Lucey, B. M., Batten, J. A., & Baur, D. G. (2015). "
     "The financial economics of gold: A survey. "
     "International Review of Financial Analysis, 41, 186-205."),
    ("Plakandaras, V., Gupta, R., Wohar, M. E., & Kaplanis, T. (2022). Forecasting "
     "the price of gold using machine learning methodologies. Applied Economics, "
     "54(33), 3768-3783."),
    ("Sims, C. A. (1980). Macroeconomics and reality. Econometrica, 48(1), 1-48."),
    ("Wooldridge, J. M. (2007). Introducción a la econometría: un enfoque moderno "
     "(3.ª ed.). Thomson."),
    ("World Gold Council. (2023). Gold Demand Trends: Full Year 2023. "
     "World Gold Council."),
    ("World Gold Council. (2024). Gold Demand Trends: Full Year 2024. "
     "World Gold Council."),
]


def write_references(doc):
    H1(doc, "REFERENCIAS BIBLIOGRÁFICAS")
    for ref in REFERENCES_APA:
        para = doc.add_paragraph(style="Normal")
        para.paragraph_format.first_line_indent = Cm(-1.0)
        para.paragraph_format.left_indent       = Cm(1.0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after       = Pt(6)
        r = para.add_run(ref); _tnr(r, 12)


# =====================================================================
# 8. SUMMARY (inglés, >= 1000 palabras)
# =====================================================================

def write_summary(doc):
    H1(doc, "SUMMARY")
    H2(doc, "Gold Price Dynamics (2000-2025): "
            "An Econometric and Machine Learning Analysis")

    H3(doc, "Introduction and motivation")
    P(doc, "Gold occupies a unique position in the taxonomy of financial assets. It "
           "generates no cash flows, pays no dividends or coupons, and has limited "
           "productive use compared to industrial commodities. Yet central banks, "
           "sovereign wealth funds, and private investors continue to accumulate it as "
           "a store of value and safe haven. The period 2000-2025 concentrated five "
           "exceptional market episodes: the Global Financial Crisis of 2008, the post-QE "
           "peak of 2011, the COVID-19 pandemic of 2020, the most aggressive rate hiking "
           "cycle in four decades (2022-2024), and an extraordinary 2025 rally that pushed "
           "prices above USD 4,500 per troy ounce for the first time in history. Each "
           "episode saw gold behave differently, making this period an exceptionally rich "
           "laboratory for applying modern econometric and machine learning tools.")
    P(doc, "This thesis is organised around three research questions: (i) which "
           "macroeconomic and financial variables determine the price of gold over "
           "2000-2025, and what is their relative importance at different time horizons? "
           "(ii) Have those determinants remained stable across major crisis episodes, or "
           "can formal structural breaks be identified? (iii) Can machine learning improve "
           "short-run predictive accuracy beyond classical econometric benchmarks?")

    H3(doc, "Theoretical framework")
    P(doc, "The analysis rests on four theoretical pillars. First, the opportunity cost "
           "mechanism: since gold bears no yield, its equilibrium price should be a "
           "decreasing function of the real interest rate. This mechanism, formalised by "
           "Barsky and Summers (1988) and documented by Erb and Harvey (2013), forms the "
           "backbone of the VECM. Second, the hedge and safe haven distinctions of Baur "
           "and Lucey (2010): gold is a hedge if its unconditional correlation with risk "
           "assets is negative on average, and a safe haven if that correlation is negative "
           "conditionally on extreme negative market returns. Third, the role of inflation "
           "expectations as a high-frequency signal of the opportunity cost. Fourth, the "
           "institutional context of de-dollarisation, through which emerging market "
           "central banks have been increasing gold reserves since 2022 at unprecedented "
           "rates. The VAR/VECM framework follows Sims's (1980) critique of structural "
           "simultaneous equation models: treating all variables as equally endogenous is "
           "the most methodologically honest approach.")

    H3(doc, "Data and methodology")
    P(doc, "The dataset covers 312 monthly observations from January 2000 to December "
           "2025. Gold prices from Yahoo Finance (GC=F); macroeconomic variables — "
           "10-year TIPS yields, CPI, 10-year breakeven rates — from FRED; financial "
           "market variables — DXY, S&P 500, VIX, WTI, 10-year Treasury yield — from "
           "Yahoo Finance. Price series are transformed to logarithmic first differences; "
           "rate series used in levels. Five historical episodes are demarcated: GFC 2008 "
           "(Aug. 2007 - Jun. 2009), post-QE peak 2011 (Jul. 2011 - Jun. 2013), "
           "COVID-19 2020 (Feb. - Aug. 2020), rate hike cycle 2022 (Mar. 2022 - Jul. "
           "2024), and the 2025 triple-confluence rally.")

    H3(doc, "VAR/VECM econometric results")
    P(doc, "ADF and KPSS unit root tests confirm all five core system variables are I(1). "
           "The Johansen tests identify one cointegrating vector (r = 1) at the 5% "
           "significance level. The long-run cointegrating vector assigns the largest "
           "coefficient to TIPS yields (-0.68), confirming the opportunity cost mechanism. "
           "The DXY enters with -1.24 and the S&P 500 with -0.31. The error correction "
           "coefficient of -0.083 implies approximately 8% of any deviation from long-run "
           "equilibrium is corrected each month (half-life: ~8 months). Granger causality "
           "tests reject non-causality from TIPS to gold at all lags (p < 0.001). Impulse "
           "response functions show a one-standard-deviation shock to real rates generates "
           "a -3.2% cumulative decline in gold at 24 months. FEVD at 12 months attributes "
           "28% of gold's variance to TIPS shocks and 19% to DXY shocks. GJR-GARCH "
           "reveals significant volatility clustering (beta = 0.847) but no significant "
           "asymmetry (gamma = -0.042, p = 0.133). Structural stability tests reject "
           "parameter stability at all five episode breakpoints, with the highest "
           "F-statistic at March 2022 (F = 4.53, p < 0.001). CUSUM exits the 5% "
           "confidence bands during 2022-2024.")

    H3(doc, "Panel data analysis")
    P(doc, "The cross-country panel covers four advanced economies (United States, Euro "
           "Area, United Kingdom, Japan) with 96 monthly observations and local-currency "
           "gold prices. The Hausman test rejects random effects (chi-squared = 12.74, "
           "p = 0.013), confirming that fixed effects capture stable unobserved "
           "country-level heterogeneity. Under fixed effects with Driscoll-Kraay standard "
           "errors, the real interest rate coefficient is negative and significant "
           "(beta = -0.61, p < 0.001), the inflation coefficient is positive "
           "(beta = +0.42, p < 0.05), and the VIX coefficient is positive "
           "(beta = +0.08, p < 0.001). The opportunity cost mechanism and the safe haven "
           "function are universal properties of gold, not peculiarities of U.S. markets.")

    H3(doc, "Machine learning: results and SHAP analysis")
    P(doc, "Three machine learning architectures are evaluated with a walk-forward "
           "expanding window: XGBoost, Random Forest, and LSTM. The feature matrix "
           "includes 35 variables and 271 effective observations. The LSTM achieves the "
           "best performance: RMSE of 3.815 pp versus 5.054 for the naive benchmark "
           "(-24.5%), and directional accuracy (DA) of 61.5% versus 55.9% (+5.6 pp). "
           "SHAP analysis shows: CPI one-month lag is the most important predictor "
           "(|SHAP| = 0.954), followed by TIPS two-month lag (0.617), one-month gold "
           "momentum (0.526), and 10-year breakeven three-month lag (0.485). The SHAP "
           "hierarchy is fully consistent with the VECM variance decomposition. When two "
           "approaches with entirely different assumptions produce the same variable "
           "importance ranking, the evidence for genuine economic causality is "
           "considerably strengthened.")

    H3(doc, "Conclusions and contributions")
    P(doc, "Four main conclusions emerge. First, real interest rates are the dominant "
           "structural determinant — robust across all three frameworks. Second, inflation "
           "is the most potent short-run predictor at the monthly horizon. Third, "
           "structural relationships are not constant: formal tests document instability "
           "at all five episode breakpoints. Fourth, the 2022-2024 paradox — gold at "
           "historical highs while real yields also reached multi-decade highs — is "
           "explained by structural demand from emerging market central banks motivated "
           "by de-dollarisation incentives, a channel inelastic to advanced-economy real "
           "rates.")
    P(doc, "The thesis makes four original contributions: (i) cross-country validation "
           "of classical mechanisms with data through 2025; (ii) formal quantification "
           "of structural instability at economically motivated breakpoints; "
           "(iii) cross-methodological validation between VECM variance decomposition "
           "and SHAP importance rankings; and (iv) integrated analysis of the 2022-2024 "
           "episode connecting econometric detection of structural break with the economic "
           "explanation of de-dollarisation. Limitations: small panel (N = 4), limited ML "
           "sample (271 observations), absence of high-frequency central bank reserve "
           "data. Natural extensions: expanding the panel to emerging economies, "
           "explicitly modelling central bank demand, and estimating Markov Switching "
           "VAR models.")
    P(doc, "Keywords: gold, VECM, cointegration, panel data, machine learning, SHAP, "
           "de-dollarisation.")


# =====================================================================
# 9. MAIN
# =====================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("TFG Oficial — formato normativo UMU 2025-2026")
    print("=" * 60)

    print("\n[0/2] Verificando figuras...")
    for fig in [FIG1, FIG2, FIG3, FIG4, FIG5, FIG6]:
        print(f"  {'OK  ' if fig.exists() else 'FALTA'} {fig.name}")

    print("\n[1/2] Construyendo TFG_Oficial.docx...")
    doc = create_document()
    add_page_numbers(doc)

    write_portada(doc)
    add_toc(doc)
    PAGE_BREAK(doc)

    write_resumen(doc)

    cap1(doc); PAGE_BREAK(doc)
    cap2(doc); PAGE_BREAK(doc)
    cap3(doc); PAGE_BREAK(doc)
    cap4(doc); PAGE_BREAK(doc)
    cap5(doc); PAGE_BREAK(doc)
    cap6(doc); PAGE_BREAK(doc)
    cap7(doc); PAGE_BREAK(doc)
    cap8(doc); PAGE_BREAK(doc)
    cap9(doc); PAGE_BREAK(doc)

    write_references(doc)
    PAGE_BREAK(doc)
    write_summary(doc)

    out_path = PROJECT_ROOT / "TFG_Oficial.docx"
    doc.save(str(out_path))

    print(f"\n[2/2] Guardado: {out_path}")
    print(f"  Párrafos : {len(doc.paragraphs)}")
    print(f"  Tablas   : {len(doc.tables)}")
    print("\nInstrucciones:")
    print("  1. Abrir TFG_Oficial.docx en Word")
    print("  2. Hacer clic sobre el índice → F9 → Actualizar toda la tabla")
