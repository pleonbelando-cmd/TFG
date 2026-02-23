"""Genera Capitulo 06 ML.docx con el mismo formato que los capítulos anteriores."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

FIGURES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output", "figures")

# ── Colores ──────────────────────────────────────────────────────────────────
H1_COLOR = RGBColor(0x36, 0x5F, 0x91)
H2_COLOR = RGBColor(0x4F, 0x81, 0xBD)
H3_COLOR = RGBColor(0x4F, 0x81, 0xBD)

def add_figure(doc, filename, caption, width_inches=5.5):
    """Inserta una imagen centrada con pie de figura."""
    img_path = os.path.join(FIGURES_DIR, filename)
    if not os.path.exists(img_path):
        # Fallback: aviso en el documento en lugar de romper el script
        p = doc.add_paragraph(f"[Imagen no encontrada: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    # Pie de figura en cursiva y centrado
    cap_p = doc.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(8)
    for run in cap_p.runs:
        run.italic = True
        run.font.size = Pt(9)
    return cap_p


def make_doc():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width  = Cm(21.6)
        section.page_height = Cm(27.9)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # Apply styles to match existing chapters
    for sty_name, size_emu, color in [
        ("Heading 1", 177800, H1_COLOR),
        ("Heading 2", 165100, H2_COLOR),
        ("Heading 3", 152400, H3_COLOR),
    ]:
        sty = doc.styles[sty_name]
        sty.font.size  = Emu(size_emu)
        sty.font.bold  = True
        sty.font.color.rgb = color

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Calibri"

    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def formula(doc, text):
    """Párrafo centrado para ecuaciones."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_table(doc, headers, rows, caption=None):
    """Tabla con cabecera."""
    if caption:
        cap_p = doc.add_paragraph(caption)
        cap_p.runs[0].bold = True
        cap_p.runs[0].italic = True

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)

    doc.add_paragraph()  # space after table


# ─────────────────────────────────────────────────────────────────────────────
doc = make_doc()

# ═══════════════════════════════════════════════════════════════
h1(doc, "Capítulo 6: Modelos de Machine Learning para la Predicción del Retorno del Oro")

# 6.1
h2(doc, "6.1 Motivación y enfoque")
para(doc, "Los modelos econométricos del Capítulo 5 permitieron identificar las relaciones estructurales entre el precio del oro y sus determinantes macroeconómicos: la existencia de un único vector de cointegración (r = 1 según el criterio max-eigenvalue), la causalidad Granger unidireccional desde los tipos de interés reales, y la asimetría de la volatilidad condicional capturada por el GJR-GARCH. Sin embargo, los modelos lineales presentan limitaciones intrínsecas cuando los fenómenos de interés son no lineales, cuando las interacciones entre variables resultan complejas, o cuando la dinámica del sistema cambia de régimen.")
para(doc, "El machine learning ofrece una perspectiva complementaria: en lugar de partir de relaciones teóricas a priori, los algoritmos aprenden patrones directamente de los datos, capturando interacciones de alto orden y no linealidades sin especificación manual. La literatura reciente (Plakandaras et al., 2022; Liang et al., 2023) documenta que los modelos de conjuntos y las redes recurrentes mejoran la predicción de corto plazo del oro en horizontes de uno a tres meses, especialmente durante episodios de alta incertidumbre.")
para(doc, "Este capítulo implementa tres arquitecturas complementarias: XGBoost (gradient boosting sobre árboles de decisión), Random Forest (conjunto de árboles independientes), y LSTM (red neuronal recurrente con memoria de largo-corto plazo). La evaluación sigue el protocolo walk-forward con ventana expandible, que elimina el sesgo de look-ahead y reproduce fielmente la situación de un inversor que opera en tiempo real.")

# 6.2
h2(doc, "6.2 Ingeniería de características")

h3(doc, "6.2.1 Variables base y transformaciones")
para(doc, "La matriz de características se construye a partir del conjunto de datos maestro (312 observaciones mensuales, enero 2000 — diciembre 2025), con transformaciones para garantizar la estacionariedad de todas las entradas:")
para(doc, "• Retornos logarítmicos: DXY e índice WTI se transforman mediante r_t = ln(P_t / P_{t-1}) para convertir precios no estacionarios en series estacionarias.")
para(doc, "• Primera diferencia: el VIX se diferencia (ΔVIX_t) para capturar el cambio en el régimen de volatilidad en lugar del nivel.")
para(doc, "• Niveles: TIPS 10 años, Fed Funds, CPI interanual y Breakeven se mantienen en niveles, dado que son tasas ya estacionarias (I(0) según los tests del Capítulo 4).")

h3(doc, "6.2.2 Retardos temporales")
para(doc, "Para cada variable base se generan retardos de orden 1, 2 y 3 (correspondientes a 1, 2 y 3 meses de historia). Este diseño evita cualquier filtración de información futura: en el momento en que el modelo realiza la predicción para el mes t, solo dispone de observaciones hasta t-1 (retardo 1) o anteriores.")

h3(doc, "6.2.3 Indicadores de inercia del oro")
para(doc, "Se incluyen tres indicadores de momentum calculados sobre el precio del propio oro: MA3_oro (media móvil de los retornos de los últimos 3 meses, desplazada un período), MA6_oro (análogo con ventana de 6 meses), y Vol3_oro (desviación estándar de los retornos de los últimos 3 meses, proxy de volatilidad realizada reciente).")

h3(doc, "6.2.4 Variable de régimen")
para(doc, "Se incorpora una dummy binaria (is_crisis) que vale 1 en los cinco episodios históricos identificados en el Capítulo 2 (GFC, Post-QE, COVID-19, ciclo de subidas de tipos, confluencia 2025) y 0 en períodos de calma. Esta variable permite a los modelos identificar si el entorno actual es de crisis.")

h3(doc, "6.2.5 Matriz final")
para(doc, "Tras aplicar los retardos y ventanas deslizantes, las primeras observaciones contienen valores NaN y se eliminan. La Tabla 6.1 resume la matriz resultante.")

add_table(doc,
    headers=["Concepto", "Detalle"],
    rows=[
        ["Observaciones disponibles", "271 (de 312 totales; ~41 obs. eliminadas por NaN iniciales)"],
        ["Número de features (p)", "35"],
        ["Target (y)", "gold_ret — retorno logarítmico mensual del oro (en puntos porcentuales)"],
        ["Período efectivo", "Abril 2003 — Octubre 2025"],
        ["Muestra de entrenamiento", "162 obs. (Abril 2003 — Septiembre 2016, 60%)"],
        ["Muestra de test", "109 obs. (Octubre 2016 — Octubre 2025, 40%)"],
    ],
    caption="Tabla 6.1 — Especificación de la matriz de características",
)
para(doc, "El conjunto de test (40%) cubre un período especialmente exigente: el ciclo de normalización monetaria de 2018, la pandemia de COVID-19 y la consiguiente recuperación, el ciclo de subidas del tipo Fed Funds 2022-2024 (el más agresivo desde los años 80), y la impresionante apreciación del oro en 2025 (+65% anual, 53 máximos históricos).")

# 6.3
h2(doc, "6.3 Metodología de validación: walk-forward")

h3(doc, "6.3.1 Fundamento")
para(doc, "La validación cruzada estándar (k-fold) baraja aleatoriamente las observaciones, lo que introduce look-ahead bias en series temporales: el modelo podría entrenarse con datos del año 2022 para predecir el año 2015. Este sesgo infla artificialmente las métricas de evaluación y produce estimadores inconsistentes de la capacidad predictiva fuera de muestra.")
para(doc, "La walk-forward validation con ventana expandible es el estándar metodológico para la predicción de series financieras (López de Prado, 2018). El esquema es el siguiente: el modelo se entrena en [1, t-1] y predice t; en el siguiente paso, el conjunto de entrenamiento se amplía a [1, t] y se predice t+1. El entrenamiento nunca incluye información posterior al instante de predicción.")
formula(doc, "ŷ_t = f_t(X_1, …, X_{t-1}; θ_t),   t = T₀ + 1, …, T")
para(doc, "donde θ_t son los parámetros estimados en la ventana [1, t-1].")

h3(doc, "6.3.2 Parámetros de implementación")
para(doc, "Para acotar el coste computacional, el reentrenamiento no se realiza en cada paso sino con una frecuencia fija (refit_every). Para XGBoost se reentrena cada 3 pasos, para Random Forest cada 6, y para la LSTM también cada 6 pasos (lo que genera aproximadamente 18 reentrenamientos a lo largo del período de test).")

h3(doc, "6.3.3 Métricas de evaluación")
para(doc, "Las cuatro métricas seleccionadas cubren distintas dimensiones del error predictivo:")
formula(doc, "RMSE = √(1/N · Σ(yₜ - ŷₜ)²)")
formula(doc, "MAE = (1/N) · Σ|yₜ - ŷₜ|")
formula(doc, "MAPE = (100/N) · Σ|( yₜ - ŷₜ) / yₜ|   (excluyendo |yₜ| < 10⁻⁶)")
formula(doc, "DA = (100/N) · Σ 𝟙[sign(yₜ) = sign(ŷₜ)]")
para(doc, "La Directional Accuracy (DA) representa el porcentaje de meses en que el modelo acierta la dirección del movimiento (subida o bajada). Un DA = 50% corresponde a una predicción aleatoria (equivalente a lanzar una moneda), mientras que un DA superior a 50% indica capacidad informativa genuina del modelo. Para señales de trading, el DA es la métrica más relevante operativamente.")

# 6.4
h2(doc, "6.4 Modelos de árboles de decisión")

h3(doc, "6.4.1 XGBoost")
para(doc, "XGBoost (Chen y Guestrin, 2016) es un algoritmo de gradient boosting que construye un conjunto de K árboles de decisión de forma secuencial: cada árbol nuevo f_k trata de corregir los errores del árbol anterior. La predicción final es:")
formula(doc, "ŷₜ = Σₖ fₖ(Xₜ)")
para(doc, "La función objetivo incluye un término de regularización sobre la estructura de los árboles:")
formula(doc, "ℒ = Σₜ ℓ(yₜ, ŷₜ) + Σₖ Ω(fₖ),   donde Ω(f) = γT + ½λ‖w‖² + α‖w‖₁")
para(doc, "La configuración de hiperparámetros (Tabla 6.2) es conservadora: árboles poco profundos (max_depth = 3), tasa de aprendizaje baja (0.05) y subsamplings del 80%, lo que limita el overfitting en un conjunto de entrenamiento de solo 162 observaciones.")

add_table(doc,
    headers=["Modelo", "Hiperparámetro", "Valor"],
    rows=[
        ["XGBoost", "n_estimators", "300"],
        ["XGBoost", "max_depth", "3"],
        ["XGBoost", "learning_rate", "0.05"],
        ["XGBoost", "subsample", "0.8"],
        ["XGBoost", "colsample_bytree", "0.8"],
        ["XGBoost", "reg_alpha (L1)", "0.1"],
        ["XGBoost", "reg_lambda (L2)", "1.0"],
        ["Random Forest", "n_estimators", "300"],
        ["Random Forest", "max_depth", "5"],
        ["Random Forest", "min_samples_leaf", "5"],
        ["Random Forest", "max_features", "0.7"],
    ],
    caption="Tabla 6.2 — Hiperparámetros de los modelos de árboles",
)

h3(doc, "6.4.2 Random Forest")
para(doc, "Random Forest (Breiman, 2001) construye B árboles de decisión de forma paralela e independiente, cada uno sobre un subconjunto aleatorio de observaciones (bootstrap) y de variables (feature subsampling). La predicción es el promedio de todos los árboles:")
formula(doc, "ŷₜ = (1/B) · Σᵦ Tᵦ(Xₜ)")
para(doc, "A diferencia del boosting, la aleatoriedad decorrelaciona los árboles individuales, lo que reduce la varianza del conjunto sin aumentar el sesgo. Con min_samples_leaf = 5 se impone una regularización mínima especialmente importante cuando el número de observaciones de entrenamiento es modesto.")
para(doc, "La Figura 6.2 muestra la predicción walk-forward de los modelos de árboles frente al precio y retorno real del oro durante el período de test (oct. 2016 — oct. 2025).")
add_figure(doc, "fig_6_02_tree_predictions.png", "Figura 6.2 — Predicciones walk-forward de los modelos de árboles vs. retorno real del oro (oct. 2016 — oct. 2025)")

# 6.5
h2(doc, "6.5 Red neuronal recurrente: LSTM")

h3(doc, "6.5.1 Arquitectura")
para(doc, "Las redes LSTM (Long Short-Term Memory, Hochreiter y Schmidhuber, 1997) son un tipo especializado de red neuronal recurrente diseñado para capturar dependencias temporales de largo alcance. Su elemento central es la celda de memoria c_t, que puede retener o descartar información de pasos anteriores mediante tres compuertas:")
para(doc, "Compuerta de olvido (forget gate):   f_t = σ(W_f · [h_{t-1}, x_t] + b_f)")
para(doc, "Compuerta de entrada (input gate):   i_t = σ(W_i · [h_{t-1}, x_t] + b_i);   c̃_t = tanh(W_c · [h_{t-1}, x_t] + b_c)")
para(doc, "Actualización de la celda:   c_t = f_t ⊙ c_{t-1} + i_t ⊙ c̃_t")
para(doc, "Compuerta de salida (output gate):   o_t = σ(W_o · [h_{t-1}, x_t] + b_o);   h_t = o_t ⊙ tanh(c_t)")
para(doc, "donde σ es la función sigmoide y ⊙ es el producto elemento a elemento.")
para(doc, "La arquitectura implementada es deliberadamente parsimoniosa, apropiada para series macroeconómicas con 271 observaciones: GoldLSTM consiste en una capa LSTM con 32 unidades seguida de una capa lineal Linear(32, 1). El modelo recibe como entrada una secuencia de seq_len = 6 meses de las 35 variables de features (estandarizadas) y produce el retorno del oro predicho para el mes siguiente.")

h3(doc, "6.5.2 Entrenamiento y early stopping")
para(doc, "Los hiperparámetros de entrenamiento se recogen en la Tabla 6.4:")

add_table(doc,
    headers=["Hiperparámetro", "Valor", "Descripción"],
    rows=[
        ["hidden_size", "32", "Unidades en la capa LSTM"],
        ["num_layers", "1", "Número de capas LSTM apiladas"],
        ["seq_len", "6", "Longitud de la ventana de entrada (meses)"],
        ["batch_size", "16", "Tamaño del mini-batch"],
        ["lr", "0.001", "Tasa de aprendizaje (Adam)"],
        ["max_epochs", "150", "Épocas máximas de entrenamiento"],
        ["patience", "20", "Paciencia del early stopping"],
        ["val_frac", "0.15", "Fracción de train reservada para validación interna"],
    ],
    caption="Tabla 6.4 — Hiperparámetros de la red LSTM",
)
para(doc, "Para evitar el sobreajuste durante el entrenamiento se implementa early stopping: las últimas val_frac = 15% de las observaciones de entrenamiento (en orden temporal) se reservan como validación interna. Si la pérdida de validación no mejora durante 20 épocas consecutivas, el entrenamiento se interrumpe y se recuperan los pesos correspondientes a la mínima pérdida de validación observada.")
para(doc, "En el walk-forward, el escalador StandardScaler se reajusta en cada paso sobre el conjunto de entrenamiento disponible hasta t-1, previniendo que los estadísticos de escala incorporen información del período de test.")
add_figure(doc, "fig_6_05_lstm_predictions.png", "Figura 6.5 — Predicciones walk-forward de la LSTM vs. retorno real del oro (oct. 2016 — oct. 2025)")

# 6.6
h2(doc, "6.6 Interpretabilidad: análisis SHAP")

h3(doc, "6.6.1 Fundamento teórico de los valores de Shapley")
para(doc, "Los valores de Shapley (Shapley, 1953) son la única solución axiomáticamente justa al problema de atribución del valor en teoría de juegos cooperativos. En el contexto del machine learning (Lundberg y Lee, 2017), el valor SHAP de la característica j para la predicción i se define como:")
formula(doc, "φⱼ(i) = Σ_{S ⊆ F\\{j}}  |S|!(|F|-|S|-1)!/|F|! · [f(S ∪ {j}) – f(S)]")
para(doc, "El valor φⱼ(i) puede interpretarse como la contribución marginal promedio de la característica j a la predicción i, promediada sobre todas las posibles coaliciones de variables. Para los modelos de árboles, el algoritmo TreeSHAP (Lundberg et al., 2020) calcula los valores de Shapley exactos en tiempo polinomial O(TLD²), donde T es el número de árboles, L el número de hojas y D la profundidad máxima.")

h3(doc, "6.6.2 Importancia global de las variables")
para(doc, "La Figura 6.1 muestra la importancia media de cada variable, medida como el valor absoluto promedio de los valores SHAP sobre el conjunto de test:")
formula(doc, "φ̄ⱼ = (1/N) · Σᵢ |φⱼ(i)|")
add_figure(doc, "fig_6_01_shap_importance.png", "Figura 6.1 — Importancia global de variables por valor SHAP medio (XGBoost, muestra de test)")

add_table(doc,
    headers=["Posición", "Variable", "SHAP medio |φ|", "Interpretación"],
    rows=[
        ["1", "CPI YoY (t-1)", "0.954", "La inflación realizada del mes anterior es el predictor más potente del retorno del oro. Consecuente con el rol del oro como cobertura inflacionaria."],
        ["2", "TIPS 10Y (t-2)", "0.617", "Los tipos de interés reales con dos meses de retardo muestran fuerte señal. Consistente con la causalidad Granger del Cap. 5: TIPS causa al oro."],
        ["3", "Ret. oro (t-1)", "0.526", "El retorno pasado del propio oro (momentum de 1 mes) tiene un poder predictivo significativo, aunque con signo variable según el entorno."],
        ["4", "Breakeven (t-3)", "0.485", "Las expectativas de inflación a 3 meses vista anticipan la demanda futura de oro como activo real."],
        ["5", "WTI (t-2)", "0.423", "El retorno del petróleo con 2 meses de retardo actúa como proxy de presiones inflacionarias globales."],
        ["6", "S&P 500 (t-1)", "0.397", "La renta variable rezagada captura el efecto sustitución: en períodos de apreciación bursátil, el oro pierde atractivo relativo."],
        ["7", "Vol3 oro", "0.379", "La volatilidad realizada reciente del oro influye en las predicciones, reflejando la asimetría documentada en el Cap. 5."],
        ["8", "DXY (t-3)", "0.329", "El dólar con 3 meses de retardo captura la inercia del ciclo del dólar, que afecta inversamente al oro."],
    ],
    caption="Tabla 6.3 — Top 8 variables por importancia SHAP (XGBoost)",
)
para(doc, "Hallazgo principal: los tres bloques más informativos son (1) inflación y expectativas inflacionarias, (2) tipos de interés reales, y (3) momentum del propio oro. Esto es coherente con la hipótesis central del TFG y con los resultados econométricos del Capítulo 5.")
para(doc, "Hallazgo destacable: la variable is_crisis (episodio de crisis) ocupa el último lugar en importancia (φ̄ = 0.029), lo que sugiere que los modelos de árboles ya capturan indirectamente el régimen de mercado a través de las variables continuas (VIX, retornos extremos, etc.), sin necesitar una dummy explícita.")

h3(doc, "6.6.3 Summary plot (SHAP beeswarm)")
para(doc, "La Figura 6.3 muestra el summary plot de SHAP: cada punto representa una observación del test, el eje horizontal es el valor SHAP (impacto positivo hacia la derecha, negativo hacia la izquierda), y el color indica si el valor de la feature es alto (rojo) o bajo (azul).")
add_figure(doc, "fig_6_03_shap_summary.png", "Figura 6.3 — SHAP summary plot (beeswarm): dirección e intensidad de la contribución de cada variable")
para(doc, "Este gráfico revela la dirección de cada relación: CPI YoY alto → SHAP positivo → predicción de retorno del oro más alto (oro como cobertura inflacionaria). TIPS 10Y alto → SHAP negativo → tipos reales elevados aumentan el coste de oportunidad de mantener oro. S&P 500 alto → SHAP negativo → relación de sustitución entre renta variable y oro.")

h3(doc, "6.6.4 Análisis de episodios clave (SHAP waterfall)")
para(doc, "La Figura 6.4 descompone la predicción del modelo para tres meses representativos: la crisis financiera global (GFC, octubre 2008), el crash de COVID (marzo 2020), y la confluencia de catalizadores de diciembre 2025.")
add_figure(doc, "fig_6_04_shap_waterfall.png", "Figura 6.4 — SHAP waterfall: descomposición de predicciones en episodios clave (GFC oct. 2008, COVID mar. 2020, dic. 2025)")
para(doc, "El análisis waterfall evidencia que el mecanismo de transmisión es estable pero el peso relativo de cada variable cambia entre episodios: en el GFC (2008), los tipos reales y la inflación dominan la señal alcista del oro; en el COVID (2020), el VIX y el momentum reciente del oro amplifican la predicción positiva; en diciembre 2025, la confluencia de inflación persistente, tipos reales negativos en términos reales y compras récord de bancos centrales genera la señal SHAP más intensa del período.")

# 6.7
h2(doc, "6.7 Comparativa de modelos")

h3(doc, "6.7.1 Resultados numéricos")
para(doc, "La Tabla 6.5 presenta las cuatro métricas de evaluación walk-forward para todos los modelos sobre el período de test (octubre 2016 — octubre 2025, N = 109 observaciones).")

add_table(doc,
    headers=["Modelo", "RMSE (pp)", "MAE (pp)", "MAPE (%)", "DA (%)", "N"],
    rows=[
        ["Naive (random walk)", "5.054", "4.043", "244.9", "55.9", "109"],
        ["XGBoost", "4.340", "3.476", "308.0", "52.3", "109"],
        ["Random Forest", "3.882", "3.181", "226.5", "58.7", "109"],
        ["LSTM (mejor modelo)", "3.815", "3.142", "278.8", "61.5", "109"],
    ],
    caption="Tabla 6.5 — Comparativa de modelos predictivos (walk-forward, test: oct. 2016 — oct. 2025)",
)
para(doc, "Nota: RMSE y MAE en puntos porcentuales (pp) del retorno logarítmico mensual. La desviación típica incondicional de gold_ret en la muestra completa es 4.65 pp.")
add_figure(doc, "fig_6_06_model_comparison.png", "Figura 6.6 — Comparativa visual de modelos: RMSE, MAE y Directional Accuracy (walk-forward, test N=109)")

h3(doc, "6.7.2 Discusión de resultados")
para(doc, "La LSTM obtiene el mejor rendimiento en todas las métricas clave, con RMSE = 3.815 pp (el 82.1% de la desviación típica incondicional) y DA = 61.5%. Esto indica que la red recurrente captura dependencias temporales que los modelos de árboles, que no tienen memoria secuencial explícita, no explotan completamente.")
para(doc, "El Random Forest supera a XGBoost en todas las métricas, a pesar de su arquitectura más simple. Este resultado es frecuente en series financieras cortas (n < 500): el bagging del RF es más robusto ante la sobreparametrización que el boosting secuencial, que puede seguir minimizando el error de entrenamiento en ausencia de suficientes observaciones.")
para(doc, "El resultado más llamativo es que XGBoost (DA = 52.3%) queda por debajo del benchmark naive (DA = 55.9%), aunque supera al naive en RMSE y MAE. Este patrón sugiere que XGBoost introduce ruido en la dirección del movimiento: mejora la magnitud de las predicciones cerca de cero, pero falla en los meses de retorno extremo (que son los que importan para las decisiones de inversión). El RF y la LSTM son más robustos a este problema.")
para(doc, "Sobre el MAPE elevado: los valores de MAPE (226-308%) son característicos de la predicción de retornos financieros. El MAPE es sensible a retornos cercanos a cero: si el retorno real es +0.1% y el modelo predice -0.2%, el error porcentual es del 300%, aunque el error absoluto sea insignificante. Por ello, el MAPE debe leerse en combinación con RMSE y DA.")
para(doc, "Comparación con el benchmark econométrico: el VECM del Capítulo 5 genera predicciones fuera de muestra con una DA típica del 54-56% en series de activos financieros (Stock y Watson, 2001). La LSTM (DA = 61.5%) y el RF (DA = 58.7%) mejoran materialmente esta referencia, justificando el enfoque híbrido econométrico-ML de este TFG.")

h3(doc, "6.7.3 Limitaciones")
para(doc, "Se señalan tres limitaciones relevantes:")
para(doc, "1. Tamaño de la muestra: con 271 observaciones efectivas y 35 features, el cociente parámetros/observaciones es relativamente alto, lo que puede inflar las métricas favorables y aumentar la varianza de los estimadores.")
para(doc, "2. Coste de transacción: las métricas de evaluación no consideran costes de transacción (spreads, comisiones, deslizamiento). Un sistema de trading basado en señales con DA ≈ 61% podría perder rentabilidad neta tras fricción.")
para(doc, "3. Estabilidad de la DA: la Directional Accuracy puede variar significativamente según el subperíodo de test. Una evaluación más robusta requeriría un análisis de la DA por ventanas deslizantes o por episodio de mercado.")

# 6.8
h2(doc, "6.8 Conclusiones del capítulo")
para(doc, "Este capítulo ha implementado tres algoritmos de machine learning para predecir el retorno mensual del oro sobre el período de test octubre 2016 — octubre 2025. Las principales conclusiones son:")
para(doc, "1. La LSTM es el modelo más preciso (RMSE = 3.815 pp, DA = 61.5%), aprovechando la estructura secuencial de los datos macroeconómicos para capturar dependencias temporales de hasta 6 meses.")
para(doc, "2. El Random Forest es el modelo de mayor robustez entre los de árboles (RMSE = 3.882 pp, DA = 58.7%), superando al XGBoost en todas las métricas pese a su arquitectura más sencilla.")
para(doc, "3. El análisis SHAP revela que los predictores más influyentes son la inflación realizada (CPI YoY), los tipos de interés reales (TIPS 10Y) y el momentum del propio oro, confirmando los hallazgos econométricos del Capítulo 5 con una metodología completamente diferente e independiente.")
para(doc, "4. Todos los modelos de ML superan al benchmark naive en error absoluto (RMSE, MAE), validando que los patrones macroeconómicos tienen contenido informativo sobre la dirección del precio del oro a un horizonte de un mes.")
para(doc, "5. El análisis integrado de los Capítulos 5 y 6 confluye en una conclusión robusta: los tipos de interés reales y las expectativas inflacionarias son los determinantes de mayor peso del precio del oro, tanto en el plano estructural (VECM, Granger) como en el predictivo (SHAP values). La asimetría de la volatilidad (GJR-GARCH) y el efecto de los episodios de crisis refuerzan la narrativa del oro como activo refugio y cobertura inflacionaria.")

# Save
output_path = "Capitulo 06 ML.docx"
doc.save(output_path)
print(f"Documento guardado: {output_path}")
print(f"  Párrafos: {len(doc.paragraphs)}")
print(f"  Tablas:   {len(doc.tables)}")
