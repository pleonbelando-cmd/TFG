"""
create_updated_chapters.py
Genera los archivos .docx actualizados para los capítulos reestructurados del TFG.
Convierte .md a .docx usando el formato definido (Calibri 11pt, colores de encabezado, márgenes).

Capítulos generados:
  - Capitulo 01 Introduccion.docx   (actualizado)
  - Capitulo 02 Marco Teorico.docx  (actualizado)
  - Capitulo 03 Vbles modelo.docx   (actualizado)
  - Capitulo 05 Econometria.docx    (actualizado)
  - Capitulo 06 Panel Data.docx     (nuevo)
  - Capitulo 07 ML.docx             (condensado, nuevo)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).parent
FIGURES_DIR = PROJECT_ROOT / "output" / "figures"

# ── Colores ────────────────────────────────────────────────────────────────
H1_COLOR = RGBColor(0x36, 0x5F, 0x91)
H2_COLOR = RGBColor(0x4F, 0x81, 0xBD)
H3_COLOR = RGBColor(0x4F, 0x81, 0xBD)


def make_doc() -> Document:
    """Crea un documento Word con el formato del TFG."""
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.05)
        section.right_margin = Cm(3.05)

    # Configurar estilos de encabezado
    for sty_name, size_pt, color, space_before_pt in [
        ("Heading 1", 14, H1_COLOR, 24),
        ("Heading 2", 13, H2_COLOR, 10),
        ("Heading 3", 11, H3_COLOR, 10),
    ]:
        sty = doc.styles[sty_name]
        sty.font.size = Pt(size_pt)
        sty.font.bold = True
        sty.font.color.rgb = color
        sty.paragraph_format.space_before = Pt(space_before_pt)

    # Estilo Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    return doc


def set_para_format(p):
    """Aplica interlineado 1.15 y espacio tras párrafo 6pt al párrafo."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '120')  # 120 twentieths = 6pt
    pPr.append(spacing)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text):
    """Añade un párrafo de texto con formato estándar, preservando negrita **text**."""
    p = doc.add_paragraph()
    set_para_format(p)
    # Procesar negrita **...**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def add_formula(doc, text):
    """Párrafo centrado para ecuaciones/fórmulas."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table_from_md(doc, lines):
    """Convierte líneas de tabla Markdown a tabla Word."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-:| ]+\|\s*$', line):
            continue  # separador
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return

    headers = rows[0]
    data = rows[1:]

    t = doc.add_table(rows=1 + len(data), cols=len(headers))
    t.style = "Table Grid"
    # Cabecera
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
        else:
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            cell.text = ''
            cell.paragraphs[0].add_run(h).bold = True
    # Filas
    for ri, row in enumerate(data):
        for ci, val in enumerate(row[:len(headers)]):
            t.rows[ri + 1].cells[ci].text = val

    doc.add_paragraph()  # espacio tras tabla


def md_to_docx(md_path: Path, docx_path: Path):
    """Convierte un archivo .md a .docx con el formato del TFG."""
    print(f"Procesando: {md_path.name} → {docx_path.name}")
    text = md_path.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()

    doc = make_doc()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Ignorar líneas de separador horizontal
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Encabezados
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), level=3)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('# '):
            add_heading(doc, line[2:].strip(), level=1)

        # Tabla markdown
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Bloque de nota/ecuación (línea con $$ o sangría de código)
        elif line.strip().startswith('$$') or line.strip().startswith('```'):
            # Saltar bloques de código/ecuaciones LaTeX
            i += 1
            while i < len(lines) and not (
                lines[i].strip().startswith('$$') or
                lines[i].strip().startswith('```')
            ):
                i += 1

        # Líneas de referencia a figuras (>)
        elif line.strip().startswith('>'):
            # Añadir como párrafo en cursiva
            content = line.strip().lstrip('> ').strip()
            if content:
                p = doc.add_paragraph(content)
                for run in p.runs:
                    run.italic = True
                p.paragraph_format.left_indent = Cm(1)

        # Párrafos normales (no vacíos)
        elif line.strip():
            # Detectar si es ecuación inline (línea que empieza con $$)
            if line.strip().startswith('$$'):
                add_formula(doc, line.strip().replace('$$', '').strip())
            else:
                add_para(doc, line.strip())

        i += 1

    doc.save(str(docx_path))
    print(f"  Guardado: {docx_path.name} ({len(doc.paragraphs)} párrafos, {len(doc.tables)} tablas)")


# ── Mapeo de capítulos ────────────────────────────────────────────────────
CHAPTERS = [
    ("capitulo_01_introduccion.md", "Capitulo 01 Introduccion.docx"),
    ("capitulo_02_marco_teorico.md", "Capitulo 02 Marco Teorico.docx"),
    ("capitulo_03_catalizadores.md", "Capitulo 03 Vbles modelo.docx"),
    ("capitulo_05_econometria.md",   "Capitulo 05 Econometria.docx"),
    ("capitulo_06_panel.md",         "Capitulo 06 Panel Data.docx"),
    ("capitulo_07_ml.md",            "Capitulo 07 ML.docx"),
]


if __name__ == "__main__":
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    ok = 0
    errors = []
    for md_name, docx_name in CHAPTERS:
        md_path = PROJECT_ROOT / md_name
        docx_path = PROJECT_ROOT / docx_name
        if not md_path.exists():
            print(f"  AVISO: {md_name} no encontrado, saltando.")
            continue
        try:
            md_to_docx(md_path, docx_path)
            ok += 1
        except Exception as e:
            print(f"  ERROR en {md_name}: {e}")
            errors.append((md_name, str(e)))

    print(f"\n=== Resumen ===")
    print(f"Generados: {ok}/{len(CHAPTERS)} capítulos")
    if errors:
        print("Errores:")
        for name, err in errors:
            print(f"  {name}: {err}")
    else:
        print("Sin errores.")
