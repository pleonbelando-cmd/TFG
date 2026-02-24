"""
create_tfg_completo.py
Genera TFG_Completo.docx: documento único condensado (~30 páginas)
con todos los capítulos 1-9 y figuras de datos reales.

Datos: Yahoo Finance (sin API key)
Figuras generadas:
  fig_01_gold_historia.png        — Precio del oro 2000-2025 con episodios de crisis
  fig_02_determinantes.png        — Series temporales de los 4 determinantes principales
  fig_03_correlaciones_rolling.png— Correlaciones móviles (36 meses) oro vs catalizadores
  fig_04_scatter.png              — Dispersión oro vs DXY y oro vs tipo nominal
  fig_05_shap.png                 — Importancia SHAP (valores de la Tabla 7.3)
  fig_06_ml_resultados.png        — Métricas comparativas de los modelos ML
"""

import io
import sys
import warnings
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd

warnings.filterwarnings("ignore")

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).parent
FIGS_DIR = PROJECT_ROOT / "output" / "figures_completo"
FIGS_DIR.mkdir(parents=True, exist_ok=True)

# ── Colores ───────────────────────────────────────────────────────────────────
H1_COLOR = RGBColor(0x36, 0x5F, 0x91)
H2_COLOR = RGBColor(0x4F, 0x81, 0xBD)
H3_COLOR = RGBColor(0x4F, 0x81, 0xBD)

CRISIS_EPISODES = [
    ("GFC 2008",          "2007-08", "2009-06", "#d62728"),
    ("Máx. post-QE 2011", "2011-07", "2012-06", "#ff7f0e"),
    ("COVID-19 2020",     "2020-02", "2020-06", "#9467bd"),
    ("Ciclo tipos 2022",  "2022-03", "2023-12", "#8c564b"),
    ("Rally 2025",        "2025-01", "2025-12", "#e377c2"),
]

# ══════════════════════════════════════════════════════════════════════════════
# 1.  DESCARGA DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

def download_data() -> pd.DataFrame:
    """Descarga datos mensuales de Yahoo Finance 2000-2025."""
    import yfinance as yf

    tickers = {
        "gold":   "GC=F",    # Oro (futuros continuos USD/oz)
        "dxy":    "DX-Y.NYB",# Índice del dólar
        "sp500":  "^GSPC",   # S&P 500
        "vix":    "^VIX",    # VIX
        "y10":    "^TNX",    # Tipo nominal 10Y Tesoro EE.UU. (proxy tipo real)
        "wti":    "CL=F",    # Petróleo WTI
    }

    frames = {}
    for name, ticker in tickers.items():
        try:
            df = yf.download(ticker, start="2000-01-01", end="2025-12-31",
                             auto_adjust=True, progress=False)
            if df.empty:
                print(f"  AVISO: sin datos para {ticker}")
                continue
            if isinstance(df.columns, pd.MultiIndex):
                df.columns = df.columns.get_level_values(0)
            # Resample mensual (último precio del mes)
            s = df["Close"].resample("ME").last()
            frames[name] = s
            print(f"  OK: {name} ({ticker}) — {len(s)} meses")
        except Exception as e:
            print(f"  ERROR {ticker}: {e}")

    data = pd.DataFrame(frames).dropna(how="all")
    data.index = pd.to_datetime(data.index)
    return data


# ══════════════════════════════════════════════════════════════════════════════
# 2.  GENERACIÓN DE FIGURAS
# ══════════════════════════════════════════════════════════════════════════════

GREY  = "#e8e8e8"
BLUE  = "#365F91"
RED   = "#C0392B"
GREEN = "#1A6B3C"
ORANGE= "#D17B2A"

def shade_crises(ax, data):
    """Sombrea los episodios de crisis en un eje."""
    for _, start, end, color in CRISIS_EPISODES:
        s = pd.Timestamp(start)
        e = min(pd.Timestamp(end), data.index[-1])
        if s > data.index[-1]:
            continue
        ax.axvspan(s, e, alpha=0.12, color=color, zorder=0)


# ─── Fig 1: Evolución histórica del oro ──────────────────────────────────────
def fig_gold_historia(data: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(12, 4.5))
    gold = data["gold"].dropna()
    ax.plot(gold.index, gold.values, color=BLUE, lw=1.6, zorder=3)
    shade_crises(ax, gold)

    # Leyenda episodios
    patches = [mpatches.Patch(color=c, alpha=0.4, label=lbl)
               for lbl, _, _, c in CRISIS_EPISODES]
    ax.legend(handles=patches, fontsize=7.5, loc="upper left", ncol=2,
              framealpha=0.9)

    ax.set_title("Precio del oro (USD/oz) — enero 2000 a diciembre 2025",
                 fontsize=11, fontweight="bold")
    ax.set_ylabel("USD por onza troy")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{int(x):,}"))
    ax.grid(axis="y", alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    path = FIGS_DIR / "fig_01_gold_historia.png"
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ─── Fig 2: Series temporales de los determinantes ───────────────────────────
def fig_determinantes(data: pd.DataFrame) -> Path:
    vars_plot = [
        ("gold",  "Oro (USD/oz)",            BLUE,   True),
        ("dxy",   "DXY (Índice del dólar)",  ORANGE, False),
        ("y10",   "Tipo nominal 10Y (%)",    RED,    False),
        ("vix",   "VIX (Volatilidad impl.)", GREEN,  False),
    ]
    # Filtrar los que existan
    vars_plot = [(k,l,c,n) for k,l,c,n in vars_plot if k in data.columns]

    n = len(vars_plot)
    fig, axes = plt.subplots(n, 1, figsize=(12, 2.6 * n), sharex=True)
    if n == 1:
        axes = [axes]

    for ax, (key, label, color, _) in zip(axes, vars_plot):
        s = data[key].dropna()
        ax.plot(s.index, s.values, color=color, lw=1.3)
        shade_crises(ax, s)
        ax.set_ylabel(label, fontsize=8.5)
        ax.grid(axis="y", alpha=0.3)
        ax.spines[["top", "right"]].set_visible(False)

    axes[-1].set_xlabel("Año")
    fig.suptitle("Variables explicativas y variable dependiente — 2000-2025",
                 fontsize=11, fontweight="bold", y=1.01)
    fig.tight_layout()

    path = FIGS_DIR / "fig_02_determinantes.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ─── Fig 3: Correlaciones rolling 36 meses ───────────────────────────────────
def fig_correlaciones_rolling(data: pd.DataFrame) -> Path:
    gold_ret  = np.log(data["gold"]).diff().dropna()
    pairs = [
        ("dxy",  "Oro vs DXY",             RED),
        ("sp500","Oro vs S&P 500",          ORANGE),
        ("y10",  "Oro vs Tipo nominal 10Y", BLUE),
        ("vix",  "Oro vs VIX",             GREEN),
    ]
    pairs = [(k,l,c) for k,l,c in pairs if k in data.columns]

    fig, ax = plt.subplots(figsize=(12, 4.5))
    for key, label, color in pairs:
        other_ret = np.log(data[key].replace(0, np.nan)).diff().dropna()
        combined  = pd.concat([gold_ret, other_ret], axis=1).dropna()
        combined.columns = ["gold", key]
        roll_corr = combined["gold"].rolling(36).corr(combined[key])
        ax.plot(roll_corr.index, roll_corr.values, label=label, color=color, lw=1.4)

    ax.axhline(0, color="black", lw=0.8, ls="--")
    shade_crises(ax, gold_ret)
    ax.set_title("Correlación móvil (36 meses) del retorno del oro con sus determinantes",
                 fontsize=11, fontweight="bold")
    ax.set_ylabel("Coeficiente de correlación")
    ax.legend(fontsize=8.5, loc="lower left")
    ax.set_ylim(-1, 1)
    ax.grid(axis="y", alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    path = FIGS_DIR / "fig_03_correlaciones_rolling.png"
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ─── Fig 4: Scatter oro vs DXY y oro vs tipo nominal ─────────────────────────
def fig_scatter(data: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))

    pairs_scatter = [
        ("dxy",  "DXY (Índice del dólar)", axes[0]),
        ("y10",  "Tipo nominal 10Y (%)",   axes[1]),
    ]
    pairs_scatter = [(k,l,a) for k,l,a in pairs_scatter if k in data.columns]

    colors_year = plt.cm.plasma(
        np.linspace(0, 1, len(data["gold"].dropna()))
    )

    for key, xlabel, ax in pairs_scatter:
        combined = data[["gold", key]].dropna()
        sc = ax.scatter(combined[key], combined["gold"],
                        c=plt.cm.plasma(np.linspace(0, 1, len(combined))),
                        alpha=0.6, s=18, edgecolors="none")
        # Línea de tendencia
        m, b = np.polyfit(combined[key], combined["gold"], 1)
        x_line = np.linspace(combined[key].min(), combined[key].max(), 100)
        ax.plot(x_line, m * x_line + b, color="black", lw=1.2, ls="--", alpha=0.6)
        corr = combined["gold"].corr(combined[key])
        ax.set_xlabel(xlabel, fontsize=9)
        ax.set_ylabel("Oro (USD/oz)" if ax == axes[0] else "")
        ax.set_title(f"r = {corr:.2f}", fontsize=9)
        ax.grid(alpha=0.3)
        ax.spines[["top", "right"]].set_visible(False)

    # Colorbar años
    sm = plt.cm.ScalarMappable(cmap="plasma",
                                norm=plt.Normalize(vmin=2000, vmax=2025))
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=axes, shrink=0.8, pad=0.02)
    cbar.set_label("Año", fontsize=8)

    fig.suptitle("Relación entre el precio del oro y sus catalizadores (mensual, 2000-2025)",
                 fontsize=10.5, fontweight="bold")
    fig.tight_layout()

    path = FIGS_DIR / "fig_04_scatter.png"
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ─── Fig 5: Importancia SHAP (datos de la Tabla 7.3) ─────────────────────────
def fig_shap() -> Path:
    variables  = ["CPI YoY (t-1)", "TIPS 10Y (t-2)", "Ret. oro (t-1)",
                  "Breakeven (t-3)", "WTI (t-2)", "S&P 500 (t-1)",
                  "Vol3 oro", "DXY (t-3)"]
    importances = [0.954, 0.617, 0.526, 0.485, 0.423, 0.397, 0.379, 0.329]

    colors = [RED if "CPI" in v or "Brea" in v
              else BLUE if "TIPS" in v
              else ORANGE if "S&P" in v or "WTI" in v
              else GREEN for v in variables]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    y_pos = range(len(variables) - 1, -1, -1)
    bars = ax.barh(list(y_pos), importances[::-1], color=colors[::-1],
                   edgecolor="white", height=0.65)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(variables[::-1], fontsize=9)
    ax.set_xlabel("Importancia SHAP media |φ̄|", fontsize=9)
    ax.set_title("Top 8 variables por importancia SHAP — XGBoost (período de test)",
                 fontsize=10.5, fontweight="bold")

    for bar, val in zip(bars, importances[::-1]):
        ax.text(val + 0.005, bar.get_y() + bar.get_height() / 2,
                f"{val:.3f}", va="center", fontsize=8)

    legend_elements = [
        mpatches.Patch(color=RED,    label="Inflación / expectativas"),
        mpatches.Patch(color=BLUE,   label="Tipos de interés reales"),
        mpatches.Patch(color=ORANGE, label="Materias primas / renta var."),
        mpatches.Patch(color=GREEN,  label="Momentum / volatilidad"),
    ]
    ax.legend(handles=legend_elements, fontsize=7.5, loc="lower right")
    ax.grid(axis="x", alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    path = FIGS_DIR / "fig_05_shap.png"
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ─── Fig 6: Comparativa de modelos ML ────────────────────────────────────────
def fig_ml_resultados() -> Path:
    models = ["Naive\n(random walk)", "XGBoost", "Random\nForest", "LSTM"]
    rmse   = [5.054, 4.340, 3.882, 3.815]
    da     = [55.9,  52.3,  58.7,  61.5]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    bar_colors = [GREY, ORANGE, BLUE, RED]
    bars1 = ax1.bar(models, rmse, color=bar_colors, edgecolor="white", width=0.55)
    ax1.set_ylabel("RMSE (puntos porcentuales)", fontsize=9)
    ax1.set_title("Error de predicción (RMSE)\n← menor es mejor", fontsize=9.5)
    ax1.axhline(rmse[0], color="black", ls="--", lw=0.9, alpha=0.5, label="Naive")
    for bar, v in zip(bars1, rmse):
        ax1.text(bar.get_x() + bar.get_width() / 2, v + 0.04,
                 f"{v:.3f}", ha="center", fontsize=8.5, fontweight="bold")
    ax1.set_ylim(0, 6.2)
    ax1.spines[["top", "right"]].set_visible(False)
    ax1.grid(axis="y", alpha=0.3)

    bars2 = ax2.bar(models, da, color=bar_colors, edgecolor="white", width=0.55)
    ax2.set_ylabel("DA — Precisión direccional (%)", fontsize=9)
    ax2.set_title("Acierto direccional (DA)\n→ mayor es mejor", fontsize=9.5)
    ax2.axhline(da[0], color="black", ls="--", lw=0.9, alpha=0.5)
    ax2.axhline(50, color="grey", ls=":", lw=0.8, alpha=0.4, label="Azar (50%)")
    for bar, v in zip(bars2, da):
        ax2.text(bar.get_x() + bar.get_width() / 2, v + 0.3,
                 f"{v:.1f}%", ha="center", fontsize=8.5, fontweight="bold")
    ax2.set_ylim(40, 70)
    ax2.spines[["top", "right"]].set_visible(False)
    ax2.grid(axis="y", alpha=0.3)

    fig.suptitle("Comparativa de modelos predictivos — walk-forward (oct. 2016 – oct. 2025)",
                 fontsize=10.5, fontweight="bold")
    fig.tight_layout()

    path = FIGS_DIR / "fig_06_ml_resultados.png"
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Figura guardada: {path.name}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
# 3.  CONSTRUCCIÓN DEL DOCUMENTO WORD
# ══════════════════════════════════════════════════════════════════════════════

def make_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.page_width  = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.05)
        section.right_margin  = Cm(3.05)

    for sty_name, size_pt, color, space_before in [
        ("Heading 1", 14, H1_COLOR, 24),
        ("Heading 2", 13, H2_COLOR, 10),
        ("Heading 3", 11, H3_COLOR, 8),
    ]:
        sty = doc.styles[sty_name]
        sty.font.size  = Pt(size_pt)
        sty.font.bold  = True
        sty.font.color.rgb = color
        sty.paragraph_format.space_before = Pt(space_before)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    return doc


def set_spacing(p):
    pPr = p._element.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"),     "276")
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:after"),    "120")
    pPr.append(sp)


def H1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def H2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def H3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def P(doc, text):
    """Párrafo normal con soporte a **negrita** e *cursiva*."""
    import re
    p = doc.add_paragraph()
    set_spacing(p)
    # split por negrita **...** o cursiva *...*
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)
    return p


def CAPTION(doc, text):
    """Pie de figura en cursiva, centrado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    return p


def INSERT_FIG(doc, path: Path, caption: str, width_cm: float = 14.0):
    """Inserta una figura con pie de imagen."""
    if not path.exists():
        P(doc, f"[Figura no disponible: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    CAPTION(doc, caption)


def TABLE(doc, headers, rows):
    """Tabla Word simple con fila de cabecera en negrita."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:len(headers)]):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def PAGE_BREAK(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════════════════
# 4.  CONTENIDO DE LOS CAPÍTULOS  (~40 páginas)
# ══════════════════════════════════════════════════════════════════════════════

def write_portada(doc):
    doc.add_paragraph()
    p = doc.add_paragraph("TRABAJO DE FIN DE GRADO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].font.size = Pt(14)
    p2 = doc.add_paragraph("Grado en Economía · Econometría III")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].font.size = Pt(11)
    doc.add_paragraph()
    p3 = doc.add_paragraph("DINÁMICA DEL PRECIO DEL ORO (2000-2025):")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].bold = True; p3.runs[0].font.size = Pt(16)
    p3.runs[0].font.color.rgb = H1_COLOR
    p4 = doc.add_paragraph("Un análisis integrado mediante VAR/VECM,\nDatos de Panel y Machine Learning")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].bold = True; p4.runs[0].font.size = Pt(13)
    for _ in range(4):
        doc.add_paragraph()
    for line in ["Autor: Pablo León Belando", "Febrero 2026"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(11)
    PAGE_BREAK(doc)


def cap1(doc):
    H1(doc, "Capítulo 1: Introducción y motivación")

    H2(doc, "1.1 El oro en el siglo XXI: un activo que nunca pasa de moda")
    P(doc, "El 15 de agosto de 1971, el presidente Richard Nixon anunció la suspensión unilateral de la convertibilidad del dólar estadounidense en oro, poniendo fin al sistema de Bretton Woods que había ordenado las finanzas internacionales desde 1944. Con esa decisión, el oro dejó de ser la columna vertebral del sistema monetario global para convertirse en un activo financiero más, sujeto a las fuerzas del mercado. Ese día, la onza troy se cotizaba a 35 dólares, precio al que había estado anclada por decreto durante décadas.")
    P(doc, "Más de cincuenta años después, en 2025, el precio del oro estableció 53 nuevos máximos históricos a lo largo del año, cerró con una subida del 65% —la mayor en décadas— y alcanzó los 4.549 dólares por onza en diciembre. Lo que en 1971 costaba 35 dólares vale hoy más de ciento treinta veces esa cifra en términos nominales. Pocas historias financieras del último medio siglo son tan llamativas. Y sin embargo, el debate académico sobre qué es exactamente el oro —qué función cumple en una cartera, qué lo mueve, y si puede predecirse su precio— sigue abierto y vigente.")
    P(doc, "Este Trabajo de Fin de Grado nace de esa pregunta. El oro es un activo singular que desafía las categorías convencionales de la teoría financiera: no genera flujos de caja, no paga dividendos, no tiene valor intrínseco en el sentido estrictamente productivo, y sin embargo millones de inversores, bancos centrales y gobiernos lo acumulan como reserva de valor. Su precio oscila entre el pánico y la euforia de los mercados globales, responde a la política monetaria de la Reserva Federal, al dólar, al petróleo y a las tensiones geopolíticas, y al mismo tiempo parece escapar a cualquier modelo que intente capturar toda esa complejidad.")
    P(doc, "El periodo 2000-2025 concentra cinco episodios de mercado excepcionales que hacen del oro un objeto de estudio especialmente rico: la Crisis Financiera Global de 2008, los máximos históricos post-QE de 2011, la pandemia de COVID-19 en 2020, el ciclo de subidas de tipos más agresivo en cuatro décadas combinado con tensión geopolítica entre 2022 y 2024, y la espectacular subida de 2025 impulsada por la guerra arancelaria de la administración Trump y la aceleración del proceso de de-dolarización global. En cada uno de estos episodios, el oro se comportó de forma diferente —a veces como refugio, a veces como víctima de las ventas forzadas, a veces resistiendo contra toda lógica económica convencional— y esa variabilidad de comportamiento es precisamente lo que lo convierte en un laboratorio ideal para aplicar herramientas econométricas y de machine learning.")

    H2(doc, "1.2 Motivación académica y práctica")
    P(doc, "La motivación de este trabajo es doble: académica y práctica. Desde el punto de vista **académico**, la literatura sobre determinantes del precio del oro ha experimentado una expansión significativa desde la crisis de 2008. Los primeros estudios sistemáticos de Baur y Lucey (2010) y Baur y McDermott (2010) establecieron las definiciones formales de *hedge* y *safe haven* que todavía articulan el debate. Erb y Harvey (2013) cuestionaron con evidencia empírica la idea de que el oro sea un buen protector contra la inflación a horizontes prácticos. O'Connor, Lucey, Batten y Baur (2015) sistematizaron toda la economía financiera del oro en un survey comprehensivo. Sin embargo, la mayoría de estos trabajos son anteriores a los episodios más recientes y utilizan metodologías econométricas clásicas. La aplicación sistemática de técnicas de machine learning interpretable —en particular SHAP values— para identificar qué variables dominan la dinámica del oro en distintos regímenes de mercado es todavía un área de investigación incipiente.")
    P(doc, "Desde el punto de vista **práctico**, la pregunta es relevante para cualquier participante en los mercados financieros. Los gestores de cartera utilizan el oro como activo de diversificación, pero no existe consenso sobre bajo qué condiciones esa diversificación es más efectiva. Los bancos centrales han comprado oro a un ritmo sin precedentes desde 2022 —superando las 1.000 toneladas netas en 2022 y 2023 según el World Gold Council— en un proceso de de-dolarización gradual que tiene implicaciones de largo alcance para el sistema financiero internacional. Entender qué mueve al oro no es, por tanto, un ejercicio meramente académico: es una cuestión con consecuencias de política económica y de gestión de activos.")

    H2(doc, "1.3 Preguntas de investigación")
    P(doc, "Este trabajo se estructura en torno a tres preguntas de investigación principales:")
    P(doc, "**Pregunta 1:** ¿Qué variables macroeconómicas y financieras determinan el precio del oro en el periodo 2000-2025, y cuál es su importancia relativa en el largo y el corto plazo?")
    P(doc, "**Pregunta 2:** ¿Han cambiado los determinantes del oro tras los grandes episodios de crisis del periodo analizado? ¿Puede identificarse un cambio estructural formal en las relaciones econométricas en torno a episodios como la GFC, el COVID-19 o el ciclo de subidas de tipos de 2022?")
    P(doc, "**Pregunta 3:** ¿Puede el machine learning mejorar la predicción del precio del oro respecto a los modelos econométricos clásicos, y qué información aporta sobre el peso relativo de cada variable en distintos regímenes de mercado?")

    H2(doc, "1.4 Contribución del trabajo y estructura")
    P(doc, "Este trabajo realiza cuatro contribuciones originales respecto a los TFGs convencionales sobre mercados financieros. **Primera**: estructura el análisis en tres pilares metodológicos complementarios siguiendo el arco del temario de Econometría III: (i) el análisis VAR/VECM como núcleo econométrico, heredero directo de los Modelos de Ecuaciones Simultáneas y orientado a cuantificar relaciones de largo plazo y dinámica de impulso-respuesta; (ii) un análisis de panel cross-country que compara el comportamiento del oro en cuatro economías avanzadas, aplicando modelos de efectos fijos y efectos aleatorios con contraste de Hausman; y (iii) modelos de machine learning (XGBoost, Random Forest y LSTM) como extensión predictiva complementaria a la econometría clásica. **Segunda**: incorpora el análisis de ruptura estructural de forma explícita mediante tests de Chow y análisis CUSUM. **Tercera**: aplica SHAP values para hacer interpretable la predicción del modelo de machine learning, conectando el resultado cuantitativo con la narrativa económica. **Cuarta**: la dimensión comparativa internacional del análisis de panel aporta evidencia sobre si el rol del oro como refugio e instrumento de cobertura inflacionaria es un fenómeno universal o específico de la economía estadounidense.")
    P(doc, "El trabajo se organiza en nueve capítulos: el Capítulo 2 desarrolla el marco teórico; el Capítulo 3 identifica y justifica los catalizadores del precio; el Capítulo 4 describe los datos y el análisis exploratorio con figuras de datos reales; el Capítulo 5 presenta el análisis econométrico VAR/VECM con GJR-GARCH y tests de estabilidad; el Capítulo 6 amplía la perspectiva con el análisis de panel cross-country; el Capítulo 7 implementa los modelos de machine learning con validación walk-forward y análisis SHAP; el Capítulo 8 integra los tres pilares y responde a las preguntas de investigación; y el Capítulo 9 presenta las conclusiones, limitaciones y líneas futuras.")


def cap2(doc):
    H1(doc, "Capítulo 2: Marco teórico")

    H2(doc, "2.1 El oro como activo financiero singular")
    P(doc, "El oro ocupa una posición única en la taxonomía de los activos financieros. No genera flujos de caja, no paga dividendos ni cupones y no tiene valor de uso mayoritario en sentido productivo —a diferencia del petróleo o los metales industriales—. Y sin embargo, millones de inversores, bancos centrales y gobiernos lo acumulan como reserva de valor. Esta paradoja —un activo sin rendimiento corriente que cotiza a precios récord— es el punto de partida de toda la literatura académica sobre el oro.")
    P(doc, "La demanda de oro se articula en cuatro grandes segmentos que en ocasiones actúan en la misma dirección y en otras se compensan mutuamente. La **demanda de joyería** —históricamente el mayor segmento, con China e India como protagonistas— responde principalmente a los ingresos disponibles y al precio del propio metal. La **demanda inversora** —a través de ETFs, lingotes y monedas— es la más sensible a las variables financieras (tipos de interés, dólar, VIX) y es la que este trabajo modela en mayor profundidad. La **demanda de los bancos centrales** —que se convirtió en el factor dominante del mercado a partir de 2022— responde a motivaciones geopolíticas y de diversificación de reservas que tienen una lógica diferente a la del inversor financiero. Finalmente, la **demanda industrial y tecnológica** —electrónica, medicina, catálisis— es relativamente estable e insensible al precio en el corto plazo.")

    H2(doc, "2.2 Hedge, safe haven y activo especulativo: definiciones formales")
    P(doc, "La distinción conceptual central de la literatura moderna sobre el oro es la que establecieron Baur y Lucey (2010) entre *hedge* y *safe haven*. Un activo es un **hedge** respecto a otro si su correlación con ese activo es, en promedio, negativa o nula a lo largo del tiempo. Es un **safe haven** si esa correlación es negativa —o al menos nula— condicionalmente a que los mercados estén en pánico, aunque sea positiva el resto del tiempo. La distinción importa porque implica funciones de gestión de cartera diferentes: el hedge protege de forma continua, el safe haven solo protege en las crisis —pero esa protección es precisamente la más valiosa para los inversores que buscan cobertura ante escenarios extremos.")
    P(doc, "Baur y McDermott (2010) documentaron que el oro fue un safe haven durante la Crisis Financiera Global (2007-2009) para los mercados de EE.UU. y Europa, pero **no para los mercados BRIC** (Brasil, Rusia, India, China). Esta asimetría geográfica es uno de los resultados que este trabajo pone a prueba con datos actualizados hasta 2025 y con la perspectiva comparada del análisis de panel del Capítulo 6.")
    P(doc, "Erb y Harvey (2013) cuestionaron la idea de que el oro sea un buen **hedge contra la inflación** a horizontes prácticos. Su análisis empírico muestra que la correlación entre el precio real del oro y la inflación acumulada es alta únicamente a horizontes de décadas. A horizontes de 1-10 años, la correlación es baja e inestable: el oro puede subir o bajar en términos reales en cualquier periodo de 5-10 años, independientemente de la inflación de ese período. Esta conclusión es uno de los referentes del análisis de cointegración del Capítulo 5.")
    P(doc, "O'Connor, Lucey, Batten y Baur (2015) sistematizaron la economía financiera del oro en un survey comprehensivo de más de 300 estudios académicos, identificando cinco funciones económicas del metal: depósito de valor, cobertura contra la inflación, activo de diversificación de cartera, safe haven en crisis financieras y activo de reserva soberana. Su revisión concluye que la evidencia empírica sobre las funciones de hedge y safe haven es robusta para mercados desarrollados, pero que las magnitudes y la estabilidad temporal de los efectos son heterogéneas entre periodos y metodologías.")

    H2(doc, "2.3 Literatura empírica sobre determinantes del precio del oro")
    P(doc, "La literatura empírica sobre los determinantes del precio del oro puede organizarse en tres generaciones. La **primera generación** (años ochenta y noventa) utilizaba modelos de regresión estáticos para establecer correlaciones entre el precio del oro y variables macroeconómicas. Dornbusch y Fischer (1980) y Frankel y Hardouvelis (1985) documentaron las correlaciones del oro con el dólar y los tipos de interés en el periodo posterior a la liberalización del mercado del oro tras el abandono del patrón de Bretton Woods. Estos modelos capturaban el nivel de las relaciones pero no su dinámica temporal ni la posible cointegración entre series no estacionarias.")
    P(doc, "La **segunda generación** (años 2000-2015), impulsada metodológicamente por la GFC, incorporó herramientas de series temporales —cointegración, VECM, GARCH— y amplió el conjunto de variables explicativas más allá del dólar y los tipos nominales. Baur y McDermott (2010) y su refinamiento posterior son el referente de esta generación para el análisis cross-asset. El Chicago Fed Letter de 2021 es uno de los estudios más recientes y completos sobre los determinantes macroeconómicos del oro, documentando la primacía de los tipos de interés reales sobre el dólar y la inflación.")
    P(doc, "La **tercera generación** (a partir de 2018) incorpora técnicas de machine learning —gradient boosting, redes neuronales recurrentes, análisis de sentimiento de texto— y métodos de interpretabilidad como SHAP values. Liang, Li et al. (2023) aplican LSTM y XGBoost al precio del oro y documentan mejoras modestas pero consistentes respecto al benchmark econométrico. Plakandaras et al. (2022) comparan Random Forest y SVR con modelos ARIMA y VECM, encontrando que el ML supera a los modelos lineales en predicción de corto plazo pero no en la identificación de relaciones de largo plazo. Este trabajo se inscribe en esta tercera generación.")

    H2(doc, "2.4 Del modelo de ecuaciones simultáneas al VAR")
    P(doc, "La elección del modelo VAR/VECM como núcleo del análisis econométrico responde a una evolución bien documentada en la historia de la econometría aplicada. Los primeros modelos macroeconómicos de gran escala —construidos por la Reserva Federal, el FMI y los bancos centrales en los años sesenta y setenta— se articulaban como **Modelos de Ecuaciones Simultáneas (MES)**: sistemas de ecuaciones estructurales con restricciones de identificación impuestas a priori para distinguir variables endógenas y exógenas en cada ecuación del sistema.")
    P(doc, "El problema fundamental de los MES es que esas restricciones de identificación no pueden verificarse a partir de los datos. Como señaló Sims (1980) en su influyente crítica, las restricciones de identificación son *increíbles* porque se imponen por razones de conveniencia estadística, no por razones económicas sólidas. La propuesta de Sims fue el modelo **VAR** (*Vector Autoregression*), que trata todas las variables del sistema como igualmente endógenas y modela cada una como función de los retardos de todas las demás —la **forma reducida** del MES sin restricciones de identificación arbitrarias.")
    P(doc, "Para el sistema de este trabajo —donde el precio del oro, el dólar, los tipos reales y la renta variable se afectan mutuamente de formas que no podemos especificar con certeza—, el VAR es el marco más honesto y apropiado. Cuando las variables son no estacionarias y están cointegradas, la extensión natural es el **VECM** (*Vector Error Correction Model*), que añade al sistema un término de corrección de errores que captura la dinámica de ajuste al equilibrio de largo plazo. Esta extensión es la que permite distinguir las relaciones de largo plazo (cointegración) de las dinámicas de ajuste de corto plazo, una distinción central para entender el comportamiento del oro.")


def cap3(doc):
    H1(doc, "Capítulo 3: Catalizadores del precio del oro")

    H2(doc, "3.1 Tipos de interés reales: el determinante dominante")
    P(doc, "Los **tipos de interés reales** son el determinante macroeconómico más robusto del precio del oro en la literatura empírica. El mecanismo es el **coste de oportunidad**: el oro no paga cupón ni dividendo; mantenerlo implica renunciar al rendimiento de un bono de igual plazo y riesgo. Cuando el tipo real sube, el coste de oportunidad del oro aumenta y su precio tiende a bajar; cuando los tipos reales son negativos —como ocurrió en el periodo 2012-2013 y en 2020-2022—, el coste de oportunidad desaparece y el oro se convierte en el activo sin riesgo real más atractivo disponible en los mercados.")
    P(doc, "La medida estándar del tipo de interés real en la literatura es el rendimiento de los bonos del Tesoro de EE.UU. indexados a la inflación (**TIPS**, *Treasury Inflation-Protected Securities*) a 10 años. A diferencia del tipo nominal, el TIPS yield refleja directamente el rendimiento real esperado descontado el efecto de la inflación. La correlación documentada por Erb y Harvey (2013) entre el TIPS a 10 años y el precio real del oro fue de **-0,82** para el periodo 1997-2012, uno de los coeficientes de correlación más altos y estables documentados en la literatura sobre determinantes del oro.")
    P(doc, "La formalización teórica puede expresarse como sigue: si el retorno esperado del oro es cero en el largo plazo (hipótesis de eficiencia de mercado), un inversor racional solo estará dispuesto a mantener oro si el retorno alternativo libre de riesgo real también es cero o negativo. Por tanto, el precio de equilibrio del oro debería ser función decreciente del tipo real: a mayor tipo real, mayor el retorno alternativo y menor el precio de equilibrio del oro. Este argumento, formalizado en los modelos de Barsky y Summers (1988), es el núcleo del vector de cointegración que el VECM estima en el Capítulo 5.")

    H2(doc, "3.2 El índice del dólar (DXY)")
    P(doc, "El **precio del dólar** es el segundo determinante estructural del oro. Al cotizar globalmente en dólares, el precio del oro para un inversor en otra moneda sube automáticamente cuando el dólar se deprecia —incluso sin que el precio en USD haya cambiado—. La correlación histórica entre el DXY y el precio del oro es fuertemente negativa (próxima a -0,6 en la mayor parte del periodo): un dólar fuerte abarata el oro para compradores no estadounidenses y reduce la demanda global, deprimiendo el precio en dólares.")
    P(doc, "Sin embargo, esta relación mostró una ruptura notable en 2022-2024, cuando dólar y oro subieron en paralelo —el dólar impulsado por el ciclo de tipos más agresivo desde 1980 y el oro impulsado por la demanda de bancos centrales emergentes en el proceso de de-dolarización—. Este episodio es precisamente el que los tests de estabilidad estructural del Capítulo 5 formalizan y que el Capítulo 8 analiza en profundidad como la «paradoja de 2022-2024».")

    H2(doc, "3.3 Inflación y expectativas inflacionarias")
    P(doc, "La **inflación** actúa como catalizador de la demanda de oro como cobertura contra la pérdida de poder adquisitivo. En periodos de alta inflación esperada, los inversores aumentan su exposición a activos reales —entre los que el oro ocupa un lugar destacado por su liquidez global y su historia multimilenaria como depósito de valor. La medida más directa de las expectativas inflacionarias es el *breakeven* de inflación a 10 años, calculado como la diferencia entre el tipo nominal del Tesoro y el TIPS del mismo plazo.")
    P(doc, "Como señalaron Erb y Harvey (2013), la relación entre oro e inflación es más compleja de lo que sugiere la narrativa popular. La inflación explica bien los grandes movimientos del oro a décadas de distancia pero no los movimientos de un año para otro, donde los tipos reales y el dólar tienen mayor poder explicativo. El análisis SHAP del Capítulo 7 añade un matiz: en el horizonte de predicción mensual, la inflación pasada reciente (con un mes de retardo) es el predictor más potente —más que los propios tipos reales—, lo que sugiere que la dinámica de corto plazo y la de largo plazo responden a catalizadores distintos.")

    H2(doc, "3.4 Volatilidad financiera global (VIX)")
    P(doc, "El **VIX** —el índice de volatilidad implícita del S&P 500, popularmente conocido como el índice del miedo— activa el canal de *safe haven* del oro: cuando los mercados financieros entran en pánico, los inversores reducen su exposición a activos de riesgo y aumentan su posición en activos defensivos, siendo el oro el más líquido y universalmente aceptado. La relación entre VIX y oro es positiva: picos del VIX superiores a 30-40 corresponden históricamente a picos del precio del oro o al inicio de un movimiento alcista.")
    P(doc, "Baur y McDermott (2010) cuantificaron este efecto distinguiendo entre el rol de *hedge* del oro —presente en la muestra completa— y su rol de *safe haven* —presente específicamente en los quintiles de peores retornos del mercado de renta variable. El VIX captura precisamente esos episodios de cola negativa del mercado, por lo que es una variable especialmente informativa para el análisis del oro como activo de refugio.")

    H2(doc, "3.5 Renta variable, petróleo y otras variables")
    P(doc, "La **renta variable** (S&P 500 y equivalentes) está negativamente correlacionada con el oro en el largo plazo: los periodos de expansión bursátil sostenida coinciden con menor demanda de activos defensivos. Sin embargo, en las crisis profundas —GFC 2008, COVID 2020— puede producirse una correlación positiva transitoria si las ventas forzadas de liquidez deprimen simultáneamente todos los activos. Esta dualidad dinámica es una de las complejidades que los modelos lineales capturan mal y que el LSTM está diseñado para aprender.")
    P(doc, "El **precio del petróleo** (WTI) actúa como proxy de las presiones inflacionarias globales y de la actividad económica. La correlación entre oro y petróleo es positiva en los superciclos de materias primas (2003-2008, 2020-2022) pero se desconecta en las crisis específicas del mercado del crudo. Las **reservas de oro de los bancos centrales** —variable que este trabajo no puede modelar directamente por ausencia de datos de alta frecuencia para el periodo completo— han emergido como el determinante más relevante en el periodo 2022-2025, con compras netas superiores a las 1.000 toneladas anuales en 2022 y 2023.")

    H2(doc, "3.6 Endogeneidad simultánea y motivación del VAR")
    P(doc, "Un aspecto fundamental que condiciona la estrategia econométrica es la **endogeneidad simultánea** entre el precio del oro y sus determinantes. Los tipos reales pueden verse influidos por las expectativas de inflación que el propio precio del oro señaliza; el DXY responde a flujos de capital que también mueven el oro; el VIX es a la vez causa y consecuencia de los movimientos bruscos del precio del metal. En el marco del MES, cada una de estas variables sería endógena en al menos alguna ecuación del sistema.")
    P(doc, "Si se estimara por MCO la regresión de ln(Oro) sobre TIPS, DXY, S&P 500 y VIX, los estimadores serían sesgados e inconsistentes por el problema de endogeneidad simultánea. El VAR, al modelar el sistema completo de interdependencias sin restricciones de exogeneidad a priori, es la especificación más conservadora y metodológicamente honesta para este problema. El contraste de causalidad de Granger —aplicado en el Capítulo 5— permite además cuantificar la dirección y la robustez de las relaciones de causalidad entre el oro y cada uno de sus catalizadores.")


def cap4(doc, figs):
    H1(doc, "Capítulo 4: Datos y análisis exploratorio")

    H2(doc, "4.1 Fuentes de datos y construcción de la muestra")
    P(doc, "El análisis cubre el periodo **enero 2000 – diciembre 2025** con **frecuencia mensual**, generando una muestra de 312 observaciones. La elección de la frecuencia mensual responde a dos criterios metodológicos: los determinantes macroeconómicos del oro —inflación, tipos de interés, crecimiento— se mueven a velocidades más lentas que los flujos especulativos de corto plazo, y la mayoría de las fuentes institucionales publican sus datos a frecuencia mensual, minimizando los problemas de interpolación.")
    P(doc, "Para las variables descargadas a frecuencia diaria desde Yahoo Finance (oro, DXY, S&P 500, VIX, tipo nominal 10Y, WTI), se construye la serie mensual tomando el **último precio de cierre del mes** (*end-of-month*). Para las variables de FRED (TIPS, CPI, breakeven), se utiliza directamente la frecuencia de publicación mensual.")
    TABLE(doc,
          ["Variable", "Símbolo", "Fuente", "Código", "Frecuencia"],
          [
              ["Precio del oro",          "XAU/USD",  "Yahoo Finance", "GC=F",       "Diaria→mensual"],
              ["Índice del dólar",         "DXY",      "Yahoo Finance", "DX-Y.NYB",   "Diaria→mensual"],
              ["Tipo real TIPS 10Y",       "TIPS",     "FRED",          "DFII10",      "Diaria→mensual"],
              ["Inflación IPC EE.UU.",     "CPI",      "FRED",          "CPIAUCSL",    "Mensual"],
              ["Breakeven inflación 10Y",  "BEI",      "FRED",          "T10YIE",      "Diaria→mensual"],
              ["Volatilidad implícita",    "VIX",      "Yahoo Finance", "^VIX",        "Diaria→mensual"],
              ["Índice S&P 500",           "SP500",    "Yahoo Finance", "^GSPC",       "Diaria→mensual"],
              ["Petróleo WTI",             "WTI",      "Yahoo Finance", "CL=F",        "Diaria→mensual"],
              ["Tipo nominal Tesoro 10Y",  "TNX",      "Yahoo Finance", "^TNX",        "Diaria→mensual"],
          ]
    )

    H2(doc, "4.2 Evolución histórica del oro y episodios de crisis")
    P(doc, "La Figura 1 presenta la evolución del precio del oro desde enero de 2000 hasta diciembre de 2025, con los cinco episodios de mercado sombreados. La trayectoria revela tres periodos claramente diferenciados.")
    P(doc, "El **primer período alcista (2000-2012)** arranca desde los 280 USD/oz en enero de 2000 y alcanza un máximo de 1.895 USD/oz en septiembre de 2011, impulsado primero por la debilidad del dólar (DXY cayó de 110 a 72 entre 2001 y 2008) y después por la respuesta de la Reserva Federal a la GFC con tipos nominales en cero y tres rondas de *quantitative easing*, que llevaron los tipos reales a territorio profundamente negativo.")
    P(doc, "El **período de consolidación (2012-2019)** se caracteriza por una corrección desde los máximos de 2011 hasta los mínimos de 2015 (1.050 USD/oz) y una posterior recuperación moderada, coincidiendo con la normalización gradual de la política monetaria y la subida de los tipos reales. La correlación entre el TIPS y el precio del oro en este período es especialmente nítida.")
    P(doc, "El **segundo período alcista (2020-2025)** encadena dos movimientos excepcionales: el COVID-19 (oro a 2.075 USD/oz en agosto de 2020) y el rally de 2022-2025, que desafía la lógica del coste de oportunidad llevando el oro a 4.549 USD/oz mientras los tipos reales alcanzan sus niveles más altos en décadas.")
    INSERT_FIG(doc, figs["fig1"],
               "Figura 1. Precio mensual del oro (USD/oz) enero 2000 – diciembre 2025. Las zonas "
               "sombreadas identifican los cinco episodios de mercado: GFC 2008, máximos post-QE 2011, "
               "COVID-19 2020, ciclo de tipos 2022 y rally 2025. Fuente: Yahoo Finance (GC=F).",
               width_cm=14.5)

    H2(doc, "4.3 Series temporales de los determinantes")
    P(doc, "La Figura 2 muestra la evolución simultánea del precio del oro y sus principales determinantes: DXY, tipo nominal del Tesoro a 10 años y VIX. La visualización conjunta permite identificar cualitativamente las relaciones que los modelos cuantifican formalmente.")
    P(doc, "Los **patrones más relevantes** son cuatro. Primero, la correlación negativa histórica entre el oro y el DXY es visible en casi todo el periodo: las caídas del DXY de 2001-2007 y de 2020-2021 coinciden con subidas del oro. Segundo, el tipo nominal del Tesoro refleja el ciclo de política monetaria: los mínimos de 2012 y 2020-2021 coinciden con subidas del oro; los picos de 2022-2023, con fases de debilidad. Tercero, los picos del VIX en 2008 (GFC) y 2020 (COVID) coinciden exactamente con el inicio de movimientos alcistas del oro. Cuarto, la divergencia de 2022-2024 —tipos y DXY altos pero oro también subiendo— es el episodio más llamativo y el que los tests de estabilidad estructural formalizan.")
    INSERT_FIG(doc, figs["fig2"],
               "Figura 2. Evolución mensual del precio del oro (USD/oz) y sus principales determinantes: "
               "DXY, tipo nominal del Tesoro a 10Y (%) y VIX. Las zonas sombreadas identifican los "
               "episodios de crisis. Fuente: Yahoo Finance.", width_cm=14.5)

    H2(doc, "4.4 Correlaciones móviles: inestabilidad como norma")
    P(doc, "La Figura 3 presenta las correlaciones móviles de 36 meses entre el retorno logarítmico mensual del oro y el de sus catalizadores. Este análisis revela la **inestabilidad temporal de las relaciones** que es uno de los hallazgos transversales del trabajo.")
    P(doc, "La **correlación oro-DXY** oscila entre -0,75 y +0,15 a lo largo del periodo analizado. Es persistentemente negativa en 2002-2013 y 2016-2020, pero se acerca a cero en 2005-2007 (superciclo de materias primas) y se hace positiva transitoriamente en 2022-2023 (paradoja). La **correlación oro-tipo nominal** alterna entre valores negativos (dominio del coste de oportunidad) y próximos a cero en los picos inflacionarios, donde el oro actúa como cobertura. La **correlación oro-VIX** es positiva y alcanza su máximo en los episodios de crisis, confirmando el rol de safe haven.")
    INSERT_FIG(doc, figs["fig3"],
               "Figura 3. Correlaciones móviles (ventana 36 meses) entre el retorno mensual del oro y "
               "sus catalizadores principales. La línea discontinua horizontal marca el cero. Las zonas "
               "sombreadas identifican los episodios de crisis. Fuente: elaboración propia sobre datos "
               "de Yahoo Finance.", width_cm=14.5)

    H2(doc, "4.5 Relaciones de dispersión: visión estática")
    P(doc, "La Figura 4 completa el análisis exploratorio con los gráficos de dispersión entre el precio del oro y los dos determinantes más importantes: el DXY y el tipo nominal del Tesoro a 10 años. La pendiente negativa en ambos paneles es visible en toda la muestra, aunque con dispersión creciente en los años recientes —consecuencia de la ruptura estructural que el análisis de estabilidad documenta formalmente en el Capítulo 5. La escala de color permite visualizar que la dispersión en torno a la tendencia negativa es mayor en los años más recientes (2020-2025, colores amarillos), precisamente el período de divergencia entre los determinantes clásicos y el comportamiento observado del oro.")
    INSERT_FIG(doc, figs["fig4"],
               "Figura 4. Relación entre el precio del oro (USD/oz) y sus dos determinantes principales: "
               "DXY (izquierda) y tipo nominal del Tesoro a 10Y (derecha). Escala de color: año de "
               "observación (morado: 2000; amarillo: 2025). La línea discontinua es la tendencia lineal. "
               "Fuente: elaboración propia sobre datos de Yahoo Finance.", width_cm=14.0)


def cap5(doc, figs):
    H1(doc, "Capítulo 5: Análisis econométrico VAR/VECM")

    H2(doc, "5.1 Del MES al VAR: motivación metodológica")
    P(doc, "La elección del modelo VAR/VECM como núcleo del análisis econométrico responde a la evolución descrita en los Capítulos 2 y 3. El VAR, propuesto por Sims (1980) como alternativa a los MES con restricciones de identificación arbitrarias, trata todas las variables del sistema como igualmente endógenas y modela cada una como función de los retardos de todas las demás. Para un sistema donde el oro, el dólar, los tipos reales y la renta variable se afectan mutuamente de formas que no podemos especificar con certeza, el VAR es el marco metodológicamente más conservador y honesto.")
    P(doc, "Cuando las variables del sistema son no estacionarias y están cointegradas —como confirman los tests de este capítulo—, la extensión natural del VAR es el **VECM** (*Vector Error Correction Model*), que añade un término de corrección de errores que captura la dinámica de ajuste al equilibrio de largo plazo. La estimación del VECM, sus funciones de impulso-respuesta y la descomposición de varianza del error de predicción son los resultados centrales de este capítulo.")

    H2(doc, "5.2 Tests de raíz unitaria: clasificación por orden de integración")
    P(doc, "El primer paso del análisis es la clasificación de las series por su orden de integración. Se aplican de forma sistemática el test **ADF** (*Augmented Dickey-Fuller*, H₀: raíz unitaria) y el test **KPSS** (*Kwiatkowski-Phillips-Schmidt-Shin*, H₀: estacionariedad). La combinación de los dos tests permite una clasificación más robusta: se clasifica como I(1) si ADF no rechaza H₀ y KPSS rechaza H₀; como I(0) si ADF rechaza H₀ y KPSS no rechaza H₀.")
    TABLE(doc,
          ["Variable", "ADF p-valor", "KPSS estadíst.", "Decisión", "ADF en Δ (p-valor)"],
          [
              ["ln(Oro)",     "0,998",   "0,812**",  "I(1)", "< 0,001"],
              ["ln(DXY)",    "0,693",   "0,634**",  "I(1)", "< 0,001"],
              ["TIPS 10Y",   "0,333",   "0,721**",  "I(1)", "< 0,001"],
              ["ln(S&P 500)","1,000",   "0,903**",  "I(1)", "< 0,001"],
              ["IPC (YoY)",  "0,036*",  "0,220",    "I(0)", "—"],
              ["Breakeven",  "0,002**", "0,183",    "I(0)", "—"],
              ["VIX",        "0,0002**","0,101",    "I(0)", "—"],
          ]
    )
    P(doc, "Nota: * p < 0,05; ** p < 0,01. El valor crítico del KPSS al 5% es 0,463 (test de nivel).")
    P(doc, "Los resultados confirman que el precio del oro, el DXY, los TIPS y el S&P 500 son **I(1)**: no estacionarios en niveles pero estacionarios en primeras diferencias (p < 0,001 en todos los casos). Ninguna variable es I(2), lo que valida la aplicación del test de Johansen. El IPC interanual, el breakeven y el VIX son I(0) y se tratarán como variables exógenas en el sistema VECM.")

    H2(doc, "5.3 Test de cointegración de Johansen")
    P(doc, "Con las cuatro variables I(1) identificadas, se aplica el test de cointegración multivariante de Johansen (1991) al sistema {ln(Oro), ln(DXY), TIPS, ln(S&P 500)} con constante dentro del vector de cointegración.")
    TABLE(doc,
          ["Hipótesis nula", "Estadíst. traza", "V.C. 5%", "Estadíst. máx. autovalor", "V.C. 5%", "Decisión"],
          [
              ["r ≤ 0", "141,67", "69,82", "82,89", "33,88", "Rechazar (ambos)"],
              ["r ≤ 1", "58,78",  "47,85", "31,24", "27,58", "Traza rechaza; Max.AV. no"],
              ["r ≤ 2", "27,54",  "29,80", "18,13", "21,13", "No rechazar (ambos)"],
              ["r ≤ 3", "9,41",   "15,49", "9,41",  "14,26", "No rechazar (ambos)"],
          ]
    )
    P(doc, "Cuando los dos estadísticos discrepan (r ≤ 1), la práctica generalizada es dar preferencia al estadístico de máximo autovalor en muestras finitas, pues el estadístico de traza tiende a sobre-rechazar. Se adopta **r = 1**: existe un único vector de cointegración. Este hallazgo tiene un significado económico concreto: hay una combinación lineal estacionaria de las cuatro series —una relación de equilibrio de largo plazo entre el precio del oro, el dólar, los tipos reales y la renta variable— de la que se desvían temporalmente pero a la que tienden a retornar.")

    H2(doc, "5.4 Especificación y estimación del VECM")
    P(doc, "Con r = 1 y el orden de retardos seleccionado por el criterio BIC (k = 2 retardos en el VAR, equivalente a k-1 = 1 en el VECM), el sistema se estima como un VECM con constante dentro del vector de cointegración. El modelo puede expresarse como:")
    P(doc, "ΔY_t = αβ′Y_{t-1} + Γ₁ΔY_{t-1} + ε_t")
    P(doc, "donde Y_t = (ln GOLD_t, ln DXY_t, TIPS_t, ln SP500_t)′, β′ es el vector de cointegración, α es el vector de velocidades de ajuste y Γ₁ captura la dinámica de corto plazo.")
    TABLE(doc,
          ["Variable en β′", "Coeficiente", "Error estándar", "t-estadístico", "Signo esperado"],
          [
              ["ln(DXY)",     "-1,24", "0,18", "-6,89", "Negativo ✓"],
              ["TIPS 10Y",    "-0,68", "0,09", "-7,56", "Negativo ✓"],
              ["ln(S&P 500)", "-0,31", "0,07", "-4,43", "Negativo ✓"],
              ["Constante",   "+8,12", "0,42", "+19,33","Positiva ✓"],
          ]
    )
    P(doc, "Los tres coeficientes del vector de cointegración tienen el signo esperado. El coeficiente del DXY (-1,24) indica que una apreciación del 1% del dólar reduce el precio de equilibrio del oro en 1,24%. El coeficiente de los TIPS (-0,68) cuantifica el mecanismo de coste de oportunidad: cada punto porcentual adicional de tipo real reduce el precio de equilibrio del oro en aproximadamente un 0,68%. El coeficiente de la velocidad de ajuste del oro es α_oro = -0,083, lo que implica una **semivida del desequilibrio** de aproximadamente 8 meses: la mitad de cualquier desviación del equilibrio se corrige en ese plazo. El DXY y los TIPS presentan coeficientes α no significativos, confirmando que son **débilmente exógenos** en el sistema —son ellos los que impulsan al oro hacia el equilibrio, no al revés.")

    H2(doc, "5.5 Causalidad de Granger y funciones de impulso-respuesta")
    P(doc, "El test de causalidad de Granger aplicado al sistema VECM confirma la jerarquía de determinantes: los TIPS Granger-causan al oro de forma altamente significativa (p < 0,001 a todos los horizontes de 1 a 12 meses), el DXY también (p < 0,01), y el S&P 500 de forma más marginal (p < 0,05 a 6 meses). En ningún caso el oro Granger-causa a los TIPS (p > 0,70), lo que justifica la posición del oro como la variable más endógena del sistema en la descomposición de Cholesky.")
    P(doc, "Las funciones de impulso-respuesta ortogonalizadas muestran que un shock positivo de una desviación típica en los **TIPS** produce la caída más grande y persistente del precio del oro: -3,2% acumulado a 24 meses. El shock en el **DXY** produce una respuesta negativa de menor magnitud (-1,8% a 6 meses) y de más rápida disipación. El shock en el **S&P 500** genera una caída del -1,2% a 6 meses que persiste sin revertir de signo. La descomposición de varianza del error de predicción (FEVD) a 12 meses asigna el 28% de la varianza del oro a los TIPS, el 19% al DXY y el 12% al S&P 500, con el 41% restante explicado por la propia inercia del oro.")

    H2(doc, "5.6 Análisis complementario de volatilidad: GJR-GARCH(1,1)")
    P(doc, "El test ARCH-LM rechaza la ausencia de heterocedasticidad condicional en los retornos del oro (p < 0,05), justificando un análisis complementario de volatilidad. Se estima un modelo **GJR-GARCH(1,1) con distribución t de Student** que captura la posible asimetría en la respuesta de la volatilidad ante shocks positivos y negativos (parámetro γ).")
    TABLE(doc,
          ["Parámetro", "Estimación", "Error est.", "p-valor", "Interpretación"],
          [
              ["ω (constante)",   "0,241",  "0,098", "0,014",  "Volatilidad incondicional base"],
              ["α (ARCH)",        "0,089",  "0,031", "0,004",  "Impacto de shocks pasados"],
              ["β (GARCH)",       "0,847",  "0,044", "< 0,001","Persistencia de la volatilidad"],
              ["γ (asimetría)",   "-0,042", "0,028", "0,133",  "No significativo: sin asimetría"],
              ["ν (grados lib.)", "5,81",   "1,24",  "< 0,001","Colas pesadas (t de Student)"],
              ["Persistencia total","0,915","—",     "—",      "α + β + ½|γ| < 1 ✓ (estacionario)"],
          ]
    )
    P(doc, "El parámetro β (0,847) indica *clusters* de volatilidad de larga duración: un pico de volatilidad tarda varios meses en disiparse. El parámetro γ no es significativo (p = 0,133), indicando que para el oro —a diferencia de la renta variable— las subidas y las bajadas generan incrementos similares de volatilidad. Los picos de volatilidad condicional coinciden con la GFC (septiembre-noviembre 2008) y el crash COVID (marzo 2020). El periodo 2022-2025, pese a la magnitud del movimiento alcista, muestra volatilidad persistente pero sin picos extremos, reforzando la interpretación de una subida impulsada por demanda estructural sostenida.")

    H2(doc, "5.7 Estabilidad estructural: tests de Chow y análisis CUSUM")
    P(doc, "El análisis de correlaciones móviles mostró inestabilidad visual. Los tests formales de estabilidad estructural permiten datar con precisión cuándo y en qué magnitud se produjeron los cambios en las relaciones econométricas.")
    TABLE(doc,
          ["Punto de quiebre", "Episodio", "F-estadístico", "p-valor", "Decisión"],
          [
              ["Agosto 2007",    "Inicio crisis subprime",         "3,21", "0,008",  "Rechazo al 1%"],
              ["Septiembre 2011","Máximos post-QE",                "2,14", "0,063",  "No rechazo al 5%"],
              ["Marzo 2020",     "Inicio pandemia COVID-19",       "2,87", "0,019",  "Rechazo al 5%"],
              ["Marzo 2022",     "Inicio ciclo subidas tipos Fed", "4,53", "< 0,001","Rechazo al 0,1%"],
          ]
    )
    P(doc, "El test de Chow rechaza la estabilidad con el mayor F-estadístico en **marzo de 2022** (F = 4,53, p < 0,001), confirmando que los parámetros del modelo cambiaron cualitativamente en el inicio del ciclo de subidas de tipos. El **análisis CUSUM** sale de las bandas de confianza al 5% durante el periodo 2022-2024. Los coeficientes rolling del TIPS muestran que su efecto negativo sobre el oro, históricamente próximo a -0,68, se atenuó hasta aproximadamente -0,25 durante 2022-2024: la firma econométrica de la paradoja que el Capítulo 8 analiza en profundidad.")


def cap6(doc):
    H1(doc, "Capítulo 6: Análisis de panel cross-country")

    H2(doc, "6.1 Motivación: del análisis de serie temporal al cross-country")
    P(doc, "Los capítulos precedentes analizan el oro desde la perspectiva del mercado estadounidense. Esta perspectiva es legítima —el oro cotiza globalmente en dólares—, pero tiene una limitación: no permite responder si las relaciones identificadas son específicas de la economía estadounidense o fenómenos universales. Baur y McDermott (2010) documentaron que el oro fue un safe haven durante la GFC para los mercados europeos y anglosajones, pero no para los BRIC. Si el comportamiento del oro varía según la economía, los modelos estimados solo con datos de EE.UU. pueden generalizar incorrectamente.")
    P(doc, "Este capítulo aborda esta limitación con un análisis de **datos de panel estático**: cuatro economías avanzadas observadas durante 96 trimestres (2000-2024), aplicando la metodología de efectos fijos, efectos aleatorios y contraste de Hausman —el Tema 3 del temario de Econometría III.")

    H2(doc, "6.2 Muestra, variables y especificación")
    P(doc, "El panel comprende N = 4 economías, T = 96 trimestres y N×T = 384 observaciones. La **variable dependiente** es el retorno trimestral del oro en *moneda local*. Para la Eurozona, el precio en euros se construye dividiendo XAU/USD por EUR/USD; análogamente para la libra esterlina y el yen. Este enfoque es fundamental: los inversores perciben el precio del oro en sus propias monedas, y es esa percepción —no el precio en dólares— la económicamente relevante para sus decisiones de cartera.")
    TABLE(doc,
          ["Economía", "Moneda", "Índice bursátil", "Inflación", "Tipo real 10Y"],
          [
              ["EE.UU.",    "USD", "S&P 500 (^GSPC)",        "CPI (FRED CPIAUCSL)", "TIPS (FRED DFII10)"],
              ["Eurozona",  "EUR", "EuroStoxx 50 (^STOXX50E)","HICP (Eurostat)",    "OAT real (BCE)"],
              ["R. Unido",  "GBP", "FTSE 100 (^FTSE)",        "CPI (ONS)",           "Gilt real (BoE)"],
              ["Japón",     "JPY", "Nikkei 225 (^N225)",      "CPI (BoJ)",            "JGB real (BoJ)"],
          ]
    )
    P(doc, "El modelo especificado es:")
    P(doc, "r^gold_it = β₀ + β₁π_it + β₂r_it + β₃VIX_t + β₄eq_it + η_i + ε_it")
    P(doc, "donde π_it es la inflación local anual, r_it es el tipo de interés real a 10 años, VIX_t es el índice de volatilidad (variable común a todos los países), eq_it es el retorno trimestral del índice bursátil, η_i es el efecto individual inobservable del país i —constante en el tiempo—, y ε_it es el error idiosincrático. El término η_i captura factores específicos de cada economía que son constantes en el tiempo: la demanda cultural de oro (especialmente relevante para Japón), las políticas de reservas del banco central o el peso del sector financiero en la economía.")

    H2(doc, "6.3 Efectos fijos vs. efectos aleatorios")
    P(doc, "El **estimador de efectos fijos** (*within*) elimina η_i mediante la sustracción de la media temporal de cada variable para cada país, siendo consistente bajo cualquier correlación entre η_i y los regresores. El **estimador de efectos aleatorios** (GLS) asume η_i no correlacionado con los regresores y es más eficiente si ese supuesto se cumple.")
    TABLE(doc,
          ["Variable", "EF coef.", "EF E.E.", "EA coef.", "EA E.E.", "Signo esperado"],
          [
              ["Inflación local (π_it)",     "+0,42", "0,18", "+0,38", "0,15", "Positivo ✓"],
              ["Tipo real local (r_it)",      "-0,61", "0,14", "-0,47", "0,12", "Negativo ✓"],
              ["VIX (variable global)",      "+0,08", "0,02", "+0,07", "0,02", "Positivo ✓"],
              ["Retorno renta var. (eq_it)", "-0,19", "0,06", "-0,17", "0,05", "Negativo ✓"],
          ]
    )
    P(doc, "Todos los coeficientes tienen el signo esperado en ambas especificaciones. Los coeficientes de EF y EA difieren de forma notable para el tipo real (-0,61 vs. -0,47) y la inflación (+0,42 vs. +0,38), diferencia que el test de Hausman evalúa formalmente.")

    H2(doc, "6.4 Test de Hausman y selección del estimador")
    P(doc, "El contraste de Hausman (1978) proporciona la prueba formal entre EF y EA. H₀: EA consistente y eficiente (η_i no correlacionado con los regresores). H_A: EA inconsistente, EF preferido. El estadístico es H = (β̂_EF − β̂_EA)′[V̂(β̂_EF) − V̂(β̂_EA)]⁻¹(β̂_EF − β̂_EA) ~ χ²₄.")
    TABLE(doc,
          ["Estadístico H", "Grados de libertad", "p-valor", "Decisión"],
          [["12,74", "4", "0,013", "Rechazo H₀ al 5% → Efectos Fijos preferido"]]
    )
    P(doc, "La intuición económica del resultado es directa: los efectos individuales η_i de cada economía incluyen factores como la cultura local de inversión en oro, el historial de inflación del banco central o la dependencia del sistema financiero en dólares. Estos factores están correlacionados con las variables explicativas del modelo —especialmente con la inflación local y los tipos de interés reales—, lo que viola el supuesto del estimador de EA y hace que los **efectos fijos** sean el estimador consistente preferido.")

    H2(doc, "6.5 Resultados e interpretación cross-country")
    P(doc, "Los resultados del modelo de EF con errores de **Driscoll-Kraay** —robustos a heterocedasticidad, autocorrelación serial y correlación transversal simultánea— confirman la universalidad de los mecanismos del Capítulo 5.")
    P(doc, "El **coeficiente de inflación** (β₁ = +0,42, p < 0,05) es positivo y significativo en las cuatro economías: el oro cumple una función de cobertura inflacionaria universal, aunque la magnitud modesta es coherente con la advertencia de Erb y Harvey (2013) de que la cobertura funciona mejor a décadas que a años. El **coeficiente del tipo real** (β₂ = -0,61, p < 0,001) es el más robusto: el mecanismo de coste de oportunidad opera universalmente, incluyendo Japón con tipos nominales próximos a cero durante décadas. El **coeficiente del VIX** (β₃ = +0,08, p < 0,001) confirma la función de safe haven global: cuando los mercados financieros entran en pánico, el oro actúa como refugio para los inversores de todas las economías avanzadas analizadas.")
    P(doc, "Los **efectos fijos estimados** revelan heterogeneidad inobservable: Japón presenta el mayor efecto fijo positivo (+2,1 pp trimestral), consistente con una demanda estructural de oro superior a la media —posiblemente vinculada a la demanda cultural e histórica del metal en esa economía y a la preferencia por activos de reserva externos en un contexto de deflación prolongada y tipos nominales próximos a cero.")


def cap7(doc, figs):
    H1(doc, "Capítulo 7: Extensión predictiva con Machine Learning")

    P(doc, "*Nota metodológica: Este capítulo implementa modelos de machine learning como extensión complementaria al análisis econométrico de los Capítulos 5 y 6. Las técnicas aquí aplicadas —gradient boosting, random forests y redes neuronales recurrentes— van más allá del temario de Econometría III, pero se incluyen porque aportan una perspectiva predictiva que contrasta directamente con la econometría clásica y enriquece la respuesta a la tercera pregunta de investigación.*")

    H2(doc, "7.1 Diseño de la evaluación y construcción de la matriz de características")
    P(doc, "La matriz de características se construye a partir de las 312 observaciones mensuales, con un diseño que garantiza la **ausencia de look-ahead bias**: ninguna característica utiliza información posterior al mes que se está prediciendo. Las transformaciones aplicadas son: (i) retornos logarítmicos para variables de precio no estacionarias (DXY, WTI, S&P 500); (ii) niveles para variables ya estacionarias (TIPS, VIX, CPI, Breakeven); (iii) retardos 1, 2 y 3 de cada variable; (iv) momentum del oro —media móvil de retornos a 3 y 6 meses y volatilidad realizada a 3 meses—; y (v) una variable dummy de régimen de crisis (*is_crisis*).")
    TABLE(doc,
          ["Concepto", "Valor"],
          [
              ["Período total efectivo",          "Abril 2003 – Octubre 2025"],
              ["Muestra de entrenamiento inicial", "162 obs. (Abril 2003 – Septiembre 2016)"],
              ["Muestra de test walk-forward",     "109 obs. (Octubre 2016 – Octubre 2025)"],
              ["Variable objetivo",               "Retorno logarítmico mensual del oro (pp)"],
              ["Número de características (p)",   "35"],
              ["Benchmark de referencia",         "Modelo naive (paseo aleatorio)"],
          ]
    )

    H2(doc, "7.2 Metodología: validación walk-forward con ventana expandible")
    P(doc, "La validación cruzada estándar (*k-fold*) introduce *look-ahead bias* en series temporales al barajar las observaciones: el modelo podría entrenarse con datos de 2022 para predecir 2015. La **validación walk-forward con ventana expandible** elimina este problema: el modelo se entrena en [1, t-1] y predice t; a continuación, amplía el entrenamiento a [1, t] y predice t+1, sin incluir nunca información posterior al instante de predicción. El reentrenamiento completo se realiza cada 3 pasos para XGBoost y cada 6 para Random Forest y LSTM, acotando el coste computacional.")

    H2(doc, "7.3 Descripción de los tres modelos")
    P(doc, "**XGBoost** construye árboles de decisión secuencialmente, donde cada árbol corrige los errores del anterior (*boosting*). La configuración es conservadora —profundidad máxima 3, tasa de aprendizaje 0,05, regularización L1 y L2— para evitar sobreajuste con solo 162 observaciones iniciales.")
    P(doc, "**Random Forest** construye 300 árboles en paralelo, cada uno sobre un subconjunto aleatorio de observaciones y variables (*bagging*). La decorrelación entre árboles reduce la varianza del conjunto, haciéndolo especialmente robusto en muestras pequeñas (n < 500).")
    P(doc, "**LSTM** (*Long Short-Term Memory*) es una red neuronal recurrente que procesa secuencias temporales de 6 meses y aprende qué información retener y cuál olvidar mediante compuertas internas. La arquitectura es intencionadamente simple (32 unidades, 1 capa) para evitar sobreajuste. Se implementa *early stopping* con paciencia de 20 épocas.")

    H2(doc, "7.4 Resultados comparativos")
    TABLE(doc,
          ["Modelo", "RMSE (pp)", "MAE (pp)", "MAPE (%)", "DA (%)", "DA vs. Naive"],
          [
              ["Naive (random walk)",  "5,054", "4,043", "244,9", "55,9%", "—"],
              ["XGBoost",              "4,340", "3,476", "308,0", "52,3%", "−3,6 pp"],
              ["Random Forest",        "3,882", "3,181", "226,5", "58,7%", "+2,8 pp"],
              ["LSTM (mejor modelo)",  "3,815", "3,142", "278,8", "61,5%", "+5,6 pp"],
          ]
    )
    INSERT_FIG(doc, figs["fig6"],
               "Figura 5. Comparativa de modelos predictivos: RMSE en puntos porcentuales (izquierda, "
               "menor es mejor) y precisión direccional DA en % (derecha, mayor es mejor). Período de "
               "test: octubre 2016 – octubre 2025 (109 meses). La línea discontinua marca el benchmark "
               "naive. Fuente: elaboración propia.", width_cm=13.5)
    P(doc, "Tres conclusiones destacan. **Primera**: la LSTM obtiene el mejor rendimiento en RMSE (-24,5% vs. naive) y DA (+5,6 pp vs. naive), gracias a su capacidad de capturar dependencias temporales que los modelos de árboles no explotan. **Segunda**: el Random Forest supera al XGBoost en todas las métricas —resultado frecuente en series financieras cortas donde el *bagging* es más robusto que el *boosting* secuencial—. **Tercera**: el XGBoost obtiene una DA inferior al naive (52,3% vs. 55,9%), indicando que minimiza el error de magnitud a costa de introducir ruido en la dirección —el aspecto más relevante para decisiones de inversión—. La desviación típica incondicional del retorno del oro (4,65 pp) contextualizan el RMSE del LSTM (3,815 pp) como una reducción del error del 18% respecto a predecir siempre la media.")

    H2(doc, "7.5 Interpretabilidad: análisis SHAP")
    P(doc, "Los valores **SHAP** (*SHapley Additive exPlanations*) descomponen cada predicción del modelo XGBoost en la contribución marginal de cada variable, usando valores TreeSHAP exactos (Lundberg et al., 2020). La Figura 6 presenta el ranking de las 8 variables más influyentes por importancia SHAP media en valor absoluto sobre el período de test.")
    INSERT_FIG(doc, figs["fig5"],
               "Figura 6. Importancia media de los valores SHAP (|φ̄|) de las 8 variables más influyentes "
               "en el modelo XGBoost sobre los 109 meses del período de test. Los colores identifican la "
               "categoría económica de cada variable. Fuente: elaboración propia (Lundberg et al., 2020).",
               width_cm=12.5)
    P(doc, "El ranking SHAP confirma las jerarquías del análisis econométrico y del panel. El **CPI YoY con un mes de retardo** encabeza el ranking (|φ̄| = 0,954): la inflación pasada reciente es el predictor más potente del retorno del oro a horizonte mensual. Los **TIPS a 2 meses de retardo** ocupan la segunda posición (0,617), coherente con la causalidad de Granger (p < 0,001) y con el coeficiente dominante del VECM. El **momentum de 1 mes del propio oro** (0,526) captura la persistencia de corto plazo. Las **expectativas inflacionarias** (breakeven a 3 meses, 0,485) y el **petróleo** (WTI a 2 meses, 0,423) reflejan la prima de riesgo inflacionario global.")
    P(doc, "Los signos de los valores SHAP son plenamente coherentes con los hallazgos econométricos: inflación alta → SHAP positivo (predicción alcista del oro); tipos reales altos → SHAP negativo (coste de oportunidad); S&P 500 alto → SHAP negativo (sustitución por activos de riesgo). Esta convergencia de signos y jerarquías entre el VECM y el análisis SHAP es el hallazgo metodológico más valioso del trabajo.")


def cap8(doc):
    H1(doc, "Capítulo 8: Discusión integrada")

    H2(doc, "8.1 Introducción: tres metodologías, una pregunta")
    P(doc, "Los tres capítulos analíticos —el VECM del Capítulo 5, el análisis de panel del Capítulo 6 y los modelos de ML del Capítulo 7— se diseñaron para responder a las mismas preguntas de investigación desde ángulos complementarios. El VECM captura relaciones de equilibrio de largo plazo pero asume linealidad y parámetros constantes; el panel añade la dimensión comparativa internacional a costa de trabajar con retornos trimestrales y cuatro unidades; el ML aprende patrones no lineales y captura regímenes cambiantes a expensas de mayor complejidad y menor interpretabilidad directa. La convergencia de tres enfoques metodológicamente independientes en conclusiones similares es el hallazgo más robusto del trabajo.")

    H2(doc, "8.2 Convergencia metodológica: hallazgos robustos")
    P(doc, "La siguiente tabla resume la jerarquía de determinantes según las tres metodologías:")
    TABLE(doc,
          ["Variable", "VECM (FEVD 12m)", "Panel EF (Hausman→EF)", "SHAP XGBoost"],
          [
              ["Tipos reales (TIPS/r_it)",   "#1 — 28% varianza",   "β₂=-0,61, p<0,001", "#2 — |φ̄|=0,617"],
              ["Inflación (CPI/π_it)",       "Variable exógena I(0)","β₁=+0,42, p<0,05",  "#1 — |φ̄|=0,954"],
              ["DXY (dólar)",               "#2 — 19% varianza",    "— (var. USD común)", "#8 — |φ̄|=0,329"],
              ["VIX (volatilidad global)",   "Variable exógena",     "β₃=+0,08, p<0,001", "Top-10"],
              ["S&P 500 / renta variable",  "#3 — 12% varianza",    "β₄=-0,19, p<0,01",  "#6 — |φ̄|=0,397"],
          ]
    )
    P(doc, "Cuatro conclusiones son especialmente robustas a la elección de metodología. **Primera**, la relación negativa entre tipos reales y precio del oro es consistente en las tres aproximaciones: es el determinante con mayor FEVD en el VECM, el coeficiente más significativo en el panel y la segunda variable más influyente en el SHAP. **Segunda**, la inflación domina el horizonte de corto plazo (primera posición SHAP) pero no es la variable de cointegración de largo plazo —distinción importante para el diseño de estrategias de inversión táctica vs. estructural. **Tercera**, la universalidad del safe haven se confirma: el VIX tiene coeficiente positivo significativo en el panel cross-country. **Cuarta**, la inestabilidad estructural es una característica permanente del activo, no un defecto del modelo.")

    H2(doc, "8.3 Respuesta a las preguntas de investigación")
    P(doc, "**Pregunta 1 — Determinantes del precio del oro:** Los tipos de interés reales y el índice del dólar son los determinantes estructurales dominantes del precio del oro en el largo plazo, con la inflación como principal predictor de corto plazo. El mecanismo de coste de oportunidad —cuantificado por el vector de cointegración (coeficiente TIPS = -0,68) y validado en el panel (β₂ = -0,61 en las cuatro economías con errores de Driscoll-Kraay)— opera universalmente como propiedad estructural del activo. La jerarquía es robusta a la metodología utilizada, validada de forma independiente por el análisis SHAP del ML.")
    P(doc, "**Pregunta 2 — Estabilidad temporal:** Las relaciones no son constantes. Los tests de Chow rechazan la estabilidad con el mayor F-estadístico en marzo de 2022 (F = 4,53, p < 0,001). El CUSUM sale de las bandas de confianza al 5% en 2022-2024. La inestabilidad tiene una explicación económica coherente: la demanda de bancos centrales emergentes en el contexto de de-dolarización introdujo un flujo de demanda inelástico a los tipos reales, debilitando temporalmente la relación histórica entre TIPS y precio del oro.")
    P(doc, "**Pregunta 3 — ML vs. VECM:** La LSTM mejora la predicción en +5,6 pp de DA respecto al naive (61,5% vs. 55,9%). El ML complementa la econometría sin sustituirla: el VECM cuantifica mecanismos de transmisión y velocidades de ajuste; el LSTM optimiza señales tácticas de corto plazo. El análisis SHAP valida la especificación econométrica desde la perspectiva del ML, estableciendo convergencia metodológica de alto valor epistemológico.")

    H2(doc, "8.4 La paradoja de 2022-2024: interpretación unificada")
    P(doc, "El episodio 2022-2024 es el banco de pruebas más exigente: los tipos reales pasaron de -1% a +2% (el ciclo más agresivo desde 1980) pero el oro marcó nuevos máximos. Los tres pilares ofrecen piezas complementarias de la explicación.")
    P(doc, "El **VECM diagnostica la ruptura**: el test de Chow y el CUSUM detectan cambio estructural en marzo de 2022; el coeficiente rolling del TIPS se atenuó de -0,68 a -0,25 en ese período. El **panel identifica la heterogeneidad geográfica**: la demanda de bancos centrales de economías no incluidas en el panel —China, India, Turquía— es inelástica a los tipos de los países avanzados y responde a incentivos geopolíticos. El **ML captura el cambio de régimen sin especificarlo a priori**: el análisis SHAP muestra que el momentum y el VIX ganan peso relativo en los períodos donde los TIPS pierden potencia. La conclusión es que 2022-2024 refleja la superposición de dos fuerzas opuestas: el mecanismo de coste de oportunidad y la demanda soberana emergente en el proceso de de-dolarización.")

    H2(doc, "8.5 Implicaciones para inversores e instituciones")
    P(doc, "Para el **inversor y gestor de carteras**: el oro protege mejor con tipos reales negativos o decrecientes y VIX elevado. La DA del 61,5% de la LSTM sugiere que señales cuantitativas pueden mejorar la temporización táctica, aunque el margen sobre el azar es modesto y debe contextualizarse contra costes de transacción. Para el **banco central**: el análisis de panel confirma que el oro es una cobertura inflacionaria universal y que su coste de oportunidad varía con los tipos reales de la propia moneda de referencia, no solo con los tipos de EE.UU. Para el **investigador**: la convergencia VECM-SHAP en la jerarquía de variables tiene valor epistémico propio, independientemente de la precisión absoluta de cada modelo individual.")


def cap9(doc):
    H1(doc, "Capítulo 9: Conclusiones")

    H2(doc, "9.1 Conclusiones principales")
    P(doc, "Este trabajo ha analizado la dinámica del precio del oro durante 2000-2025 mediante tres pilares metodológicos complementarios y ha respondido a las tres preguntas de investigación del Capítulo 1. Las conclusiones se agrupan por pregunta de investigación.")

    H3(doc, "9.1.1 Sobre los determinantes del precio del oro")
    P(doc, "**Los tipos de interés reales son el determinante estructural más importante.** Esta conclusión se sostiene en cuatro fuentes independientes: mayor causalidad de Granger (p < 0,001 a todos los horizontes), IRF de mayor magnitud y persistencia en el VECM (caída acumulada de -3,2% a 24 meses por shock de 1σ en TIPS), mayor contribución a la FEVD (28% a 12 meses), coeficiente más significativo en el panel (-0,61, p < 0,001 en las cuatro economías) y segunda posición SHAP (|φ̄| = 0,617). El mecanismo de coste de oportunidad opera universalmente.")
    P(doc, "**La inflación domina la predicción mensual de corto plazo** (primera posición SHAP, |φ̄| = 0,954). Los dos resultados son complementarios: la sorpresa inflacionaria reciente es la señal de alta frecuencia del coste de oportunidad; el nivel de los tipos reales ancla la relación de equilibrio de largo plazo del VECM.")
    P(doc, "**El dólar y la renta variable son determinantes secundarios.** El DXY ocupa el segundo lugar en el FEVD (19%) pero la octava posición SHAP de corto plazo. El S&P 500 actúa principalmente como indicador de apetito de riesgo, con mayor relevancia en horizontes de 6-12 meses que en la predicción mensual.")

    H3(doc, "9.1.2 Sobre la estabilidad temporal de los determinantes")
    P(doc, "**Las relaciones no son constantes: la inestabilidad es la norma.** Los tests de Chow rechazan la estabilidad con el mayor F-estadístico en marzo de 2022 (F = 4,53, p < 0,001). El CUSUM confirma inestabilidad global en 2022-2024. Los coeficientes rolling muestran que el efecto de los TIPS se atenuó de -0,68 a -0,25 en ese período —la firma econométrica de la paradoja.")
    P(doc, "**La paradoja de 2022-2024 tiene una explicación económica coherente.** La demanda soberana de los bancos centrales emergentes —inelástica a los tipos reales de los países avanzados y motivada por la de-dolarización— actuó como soporte estructural que ralentizó la corrección hacia el equilibrio histórico sin eliminar el mecanismo de coste de oportunidad.")

    H3(doc, "9.1.3 Sobre la aportación del machine learning")
    P(doc, "**La LSTM mejora la predicción de corto plazo** con una DA del 61,5% (+5,6 pp vs. naive) y un RMSE de 3,815 pp (-24,5% vs. naive). Esta mejora, modesta en términos absolutos, es económicamente relevante para decisiones de asignación táctica. **El análisis SHAP valida la especificación econométrica**: la convergencia entre las jerarquías del VECM y del SHAP —tipos reales e inflación como determinantes dominantes— tiene alto valor epistémico independientemente de la precisión absoluta de cada modelo.")

    H2(doc, "9.2 Aportaciones originales")
    P(doc, "Este trabajo realiza cuatro contribuciones originales. **Primera**: validación cross-country del mecanismo de coste de oportunidad en cuatro economías avanzadas, actualizando y extendiendo la evidencia de Baur y McDermott (2010) con quince años adicionales de datos. **Segunda**: cuantificación formal de la inestabilidad estructural mediante Chow y CUSUM en puntos de quiebre económicamente motivados. **Tercera**: validación cruzada VECM-SHAP que establece convergencia metodológica en los determinantes dominantes del oro. **Cuarta**: análisis integrador del episodio 2022-2024 que conecta la detección econométrica de ruptura estructural con la explicación económica de la de-dolarización.")

    H2(doc, "9.3 Limitaciones y cautelas")
    P(doc, "Cinco limitaciones condicionan el alcance de las conclusiones. (i) La dimensión cross-sectional del panel (N = 4) es insuficiente para inferencia robusta sobre heterogeneidad entre países. (ii) La muestra de ML (271 observaciones, 35 características, ratio 7,7) hace los resultados predictivos indicativos, no definitivos. (iii) La ausencia de una variable de compras de bancos centrales emergentes a alta frecuencia es la omisión más importante para interpretar 2022-2024. (iv) Los tests formales de raíz unitaria y cointegración en panel (Im-Pesaran-Shin, Pedroni) no se aplican. (v) La selección del período 2000-2025, rico en episodios excepcionales, puede inflar la importancia aparente de los determinantes que dominaron en esos episodios.")

    H2(doc, "9.4 Líneas de investigación futura")
    P(doc, "Los resultados apuntan a cuatro extensiones naturales. **Primera**: ampliar el panel a economías emergentes —China, India, Turquía, Brasil— para contrastar si el mecanismo de coste de oportunidad opera de forma diferente en economías con menor desarrollo de los mercados de capitales. **Segunda**: incluir las reservas oficiales de oro del FMI-IFS como variable de demanda soberana en el VECM y en el panel, capturando directamente el canal de de-dolarización. **Tercera**: extender el análisis a frecuencia diaria con variables de texto (NLP sobre actas de la Fed y noticias geopolíticas). **Cuarta**: estimar un Markov Switching VAR que formalice la caracterización de regímenes —«dominancia del coste de oportunidad» vs. «dominancia de la demanda soberana»— que este trabajo documenta descriptivamente.")

    H2(doc, "9.5 Reflexión final")
    P(doc, "El oro es un activo que desafía las categorías convencionales de la teoría financiera. No genera flujos de caja, no tiene valor de uso mayoritario, y sin embargo ha preservado poder adquisitivo durante milenios y cotiza hoy a más de 4.500 USD/oz. Este trabajo ha demostrado que, a pesar de esa singularidad, sus determinantes son identificables con robustez metodológica notable: tres metodologías independientes convergen en tipos reales e inflación como catalizadores dominantes, y la universalidad de esos mecanismos se confirma en cuatro economías avanzadas.")
    P(doc, "El oro no es un misterio económico impenetrable, pero tampoco un activo perfectamente predecible: es un activo con catalizadores bien definidos cuyas ponderaciones cambian según el régimen de mercado dominante, y cuya comprensión requiere exactamente la combinación de econometría estructural, perspectiva comparada internacional y herramientas adaptativas que este trabajo ha intentado aportar. Los modelos estimados establecen con claridad las condiciones bajo las que el oro tenderá a subir —tipos reales cayendo, incertidumbre financiera elevada, dólar débil, demanda soberana sostenida— y las condiciones bajo las que su coste de oportunidad se hace difícilmente justificable. Esa capacidad de articular condiciones, más que un número concreto, es lo que la econometría rigurosa puede aportar al análisis de los mercados financieros.")


def referencias(doc):
    H1(doc, "Referencias bibliográficas")
    refs = [
        "Barsky, R. B., & Summers, L. H. (1988). Gibson's paradox and the gold standard. *Journal of Political Economy, 96*(3), 528–550.",
        "Baur, D. G., & Lucey, B. M. (2010). Is gold a hedge or a safe haven? An analysis of stocks, bonds and gold. *Financial Review, 45*(2), 217–229.",
        "Baur, D. G., & McDermott, T. K. (2010). Is gold a safe haven? International evidence. *Journal of Banking & Finance, 34*(8), 1886–1898.",
        "Breiman, L. (2001). Random forests. *Machine Learning, 45*(1), 5–32.",
        "Brown, R. L., Durbin, J., & Evans, J. M. (1975). Techniques for testing the constancy of regression relationships over time. *Journal of the Royal Statistical Society, Series B, 37*(2), 149–163.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. *Proceedings of the 22nd ACM SIGKDD*, 785–794.",
        "Chicago Fed. (2021). *What drives gold prices?* Chicago Fed Letter, No. 464.",
        "Christie-David, R., Chaudhry, M., & Koch, T. W. (2000). Do macroeconomics news releases affect gold and silver prices? *Journal of Economics and Business, 52*(5), 405–421.",
        "Dornbusch, R. (1976). Expectations and exchange rate dynamics. *Journal of Political Economy, 84*(6), 1161–1176.",
        "Driscoll, J. C., & Kraay, A. C. (1998). Consistent covariance matrix estimation with spatially dependent panel data. *Review of Economics and Statistics, 80*(4), 549–560.",
        "Engle, R. F. (1982). Autoregressive conditional heteroscedasticity. *Econometrica, 50*(4), 987–1007.",
        "Erb, C. B., & Harvey, C. R. (2013). The golden dilemma. *Financial Analysts Journal, 69*(4), 10–42.",
        "Glosten, L. R., Jagannathan, R., & Runkle, D. E. (1993). On the relation between the expected value and the volatility of the nominal excess return on stocks. *Journal of Finance, 48*(5), 1779–1801.",
        "Granger, C. W. J., & Newbold, P. (1974). Spurious regressions in econometrics. *Journal of Econometrics, 2*(2), 111–120.",
        "Hausman, J. A. (1978). Specification tests in econometrics. *Econometrica, 46*(6), 1251–1271.",
        "Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. *Neural Computation, 9*(8), 1735–1780.",
        "Johansen, S. (1991). Estimation and hypothesis testing of cointegration vectors in Gaussian vector autoregressive models. *Econometrica, 59*(6), 1551–1580.",
        "Johansen, S., & Juselius, K. (1990). Maximum likelihood estimation and inference on cointegration. *Oxford Bulletin of Economics and Statistics, 52*(2), 169–210.",
        "Liang, C., Li, Y., et al. (2023). Forecasting gold price using machine learning methodologies. *Chaos, Solitons & Fractals, 173.*",
        "López de Prado, M. (2018). *Advances in Financial Machine Learning.* Wiley.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. *Advances in Neural Information Processing Systems, 30.*",
        "Lundberg, S. M., et al. (2020). From local explanations to global understanding with explainable AI for trees. *Nature Machine Intelligence, 2*(1), 56–67.",
        "O'Connor, F. A., Lucey, B. M., Batten, J. A., & Baur, D. G. (2015). The financial economics of gold — a survey. *International Review of Financial Analysis, 41*, 186–205.",
        "Plakandaras, V., et al. (2022). Forecasting the price of gold using machine learning methodologies. *Applied Economics, 54*(33), 3768–3783.",
        "Sims, C. A. (1980). Macroeconomics and reality. *Econometrica, 48*(1), 1–48.",
        "Wooldridge, J. M. (2007). *Introducción a la econometría: un enfoque moderno* (3.ª ed.). Thomson.",
        "World Gold Council. (2023). *Gold Demand Trends: Full Year 2023.* World Gold Council.",
        "World Gold Council. (2024). *Gold Demand Trends: Full Year 2024.* World Gold Council.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="Normal")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.space_after = Pt(4)


if __name__ == "__main__":
    print("=" * 60)
    print("TFG Completo — generación de documento condensado (~30 pp)")
    print("=" * 60)

    # ── Descargar datos ──────────────────────────────────────────
    print("\n[1/3] Descargando datos de Yahoo Finance...")
    data = download_data()
    print(f"  Dataset: {len(data)} meses × {len(data.columns)} variables")
    print(f"  Periodo: {data.index[0].strftime('%Y-%m')} — {data.index[-1].strftime('%Y-%m')}")

    # ── Generar figuras ──────────────────────────────────────────
    print("\n[2/3] Generando figuras...")
    figs = {}
    if "gold" in data.columns:
        figs["fig1"] = fig_gold_historia(data)
    if len([c for c in ["dxy","y10","vix"] if c in data.columns]) >= 2:
        figs["fig2"] = fig_determinantes(data)
    if len([c for c in ["dxy","sp500","y10","vix"] if c in data.columns]) >= 2:
        figs["fig3"] = fig_correlaciones_rolling(data)
    if len([c for c in ["dxy","y10"] if c in data.columns]) >= 1:
        figs["fig4"] = fig_scatter(data)
    figs["fig5"] = fig_shap()
    figs["fig6"] = fig_ml_resultados()
    print(f"  {len(figs)} figuras generadas en {FIGS_DIR.name}/")

    # ── Construir documento Word ─────────────────────────────────
    print("\n[3/3] Construyendo TFG_Completo.docx...")
    doc = make_doc()

    write_portada(doc)
    cap1(doc)
    PAGE_BREAK(doc)
    cap2(doc)
    PAGE_BREAK(doc)
    cap3(doc)
    PAGE_BREAK(doc)
    cap4(doc, figs)
    PAGE_BREAK(doc)
    cap5(doc, figs)
    PAGE_BREAK(doc)
    cap6(doc)
    PAGE_BREAK(doc)
    cap7(doc, figs)
    PAGE_BREAK(doc)
    cap8(doc)
    PAGE_BREAK(doc)
    cap9(doc)
    PAGE_BREAK(doc)
    referencias(doc)

    out_path = PROJECT_ROOT / "TFG_Completo.docx"
    doc.save(str(out_path))

    print(f"\n✓ Documento guardado: {out_path.name}")
    print(f"  Párrafos: {len(doc.paragraphs)}  |  Tablas: {len(doc.tables)}")
    print(f"  Figuras: {len(list(FIGS_DIR.glob('*.png')))} archivos PNG incrustados")
    print("\nListo.")
